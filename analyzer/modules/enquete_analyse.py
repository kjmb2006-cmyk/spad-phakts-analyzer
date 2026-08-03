"""
SPAD Analyzer — Moteur d'analyse épidémiologique générique

À partir d'un dictionnaire de variables VALIDÉ PAR L'UTILISATEUR (voir
modules/xlsform_dictionary.py pour la génération initiale, éditable ensuite
dans l'interface) et des données réelles d'une enquête Kobo/ODK quelconque,
calcule :
  - statistiques univariées par indicateur, avec IC95%
  - scores composites — regroupement ENTIÈREMENT piloté par l'utilisateur
    (colonnes `domaine`, `inclure_score_composite`, `sens_item`,
    `valeurs_favorables` du dictionnaire) : aucune détection sémantique
    automatique du contenu des questions n'est tentée ici, contrairement au
    type de traitement statistique (dérivé du xType/type XLSForm, lui,
    objectivement déterminable)
  - croisement des scores composites (corrélation + comparaison de moyennes)
  - stratification par les variables marquées role='stratification' dans le
    dictionnaire (limité aux select_one — pas de sens à stratifier par une
    variable numérique ou texte libre sans binning, non fait ici)
  - tableau de qualité des données (complétion, valeurs aberrantes,
    couverture dictionnaire ↔ données réellement collectées, alertes de
    complétion quasi nulle vs variable conditionnante bien remplie)
  - rapport Word — sections dynamiques selon ce qui a été trouvé/configuré,
    incluant en tête la complétude par groupe déjà calculée par
    modules/projets.py (réutilisée, pas recalculée ici)

Ne connaît rien à SPAD ni à un formulaire particulier.
"""
import io
import os
import re
import warnings
import numpy as np
import pandas as pd
from scipy import stats

from modules.xlsform_dictionary import _extract_suffix, _xtype_key  # noqa: F401 (réutilisé)

warnings.filterwarnings('ignore')

ALPHA = 0.05
Z = stats.norm.ppf(1 - ALPHA / 2)

KOBO_SYS_COLS = {
    'meta/instanceID', '_xform_id_string', '_attachments', '_geolocation',
    '_id', '_uuid', '_submission_time', '_validation_status', '_notes',
    '_status', '_submitted_by', '__version__', '_tags', 'formhub/uuid',
}


# ─────────────────────────────────────────────────────────────────────────
# Rapprochement dictionnaire ↔ données
# ─────────────────────────────────────────────────────────────────────────

def match_columns(dic, data):
    """Associe chaque variable du dictionnaire à sa colonne réelle dans les
    données (nom exact ou suffixe `.../nom`, pour les champs préfixés par
    un groupe Kobo). Renvoie (dic_avec_colonne, variables_non_trouvees,
    colonnes_hors_dictionnaire)."""
    dic = dic.copy()
    data_map = {}
    for c in data.columns:
        data_map.setdefault(c.split('/')[-1], c)
    dic['colonne_donnees'] = dic['nom'].map(data_map)

    reperees = dic[dic['role'].isin(['indicateur', 'stratification'])]
    manquantes = reperees[reperees['colonne_donnees'].isna()]

    dict_names = set(dic['nom'])
    extra_cols = [c for c in data.columns
                  if c.split('/')[-1] not in dict_names and c not in KOBO_SYS_COLS]
    return dic, manquantes, extra_cols


# ─────────────────────────────────────────────────────────────────────────
# Statistiques univariées
# ─────────────────────────────────────────────────────────────────────────

def wilson_ci(k, n):
    if n == 0:
        return (np.nan, np.nan)
    phat = k / n
    denom = 1 + Z**2 / n
    centre = (phat + Z**2 / (2 * n)) / denom
    marge = (Z * np.sqrt(phat * (1 - phat) / n + Z**2 / (4 * n**2))) / denom
    return (max(0, centre - marge) * 100, min(1, centre + marge) * 100)


def univariate_table(dic, data):
    rows = []
    indicateurs = dic[(dic['role'] == 'indicateur') & dic['colonne_donnees'].notna()]

    for _, ind in indicateurs.iterrows():
        col = ind['colonne_donnees']
        is_multi = str(ind['type_xlsform']).startswith('select_multiple')
        xkey = _xtype_key(ind['suffixe_phakts'] or None, is_multi)
        if xkey is None:
            type_simple = str(ind['type_xlsform']).split(' ')[0]
            xkey = {'select_one': 'X_one', 'select_multiple': 'X_multi',
                    'integer': 'period_raw', 'decimal': 'period_raw', 'range': 'period_raw'}.get(type_simple)
        serie = data[col]
        n_total = len(serie)
        n_valid = serie.notna().sum()

        if xkey in ('B', 'C', 'X_one'):
            vc = serie.dropna().value_counts()
            for modalite, k in vc.items():
                lo, hi = wilson_ci(k, n_valid)
                rows.append({
                    'variable': ind['nom'], 'section': ind['section'], 'label': ind['label'],
                    'modalite': modalite, 'n_valide': n_valid, 'n_total': n_total,
                    'effectif': k, 'pourcentage': round(100 * k / n_valid, 1) if n_valid else np.nan,
                    'ic95_inf': round(lo, 1) if not np.isnan(lo) else np.nan,
                    'ic95_sup': round(hi, 1) if not np.isnan(hi) else np.nan,
                    'statistique': 'proportion',
                })

        elif xkey == 'X_multi':
            tokens_par_ligne = serie.dropna().apply(lambda s: str(s).split())
            n_reponses = len(tokens_par_ligne)
            all_tokens = sorted(set(t for toks in tokens_par_ligne for t in toks))
            for tok in all_tokens:
                k = sum(1 for toks in tokens_par_ligne if tok in toks)
                lo, hi = wilson_ci(k, n_reponses)
                rows.append({
                    'variable': ind['nom'], 'section': ind['section'], 'label': ind['label'],
                    'modalite': tok, 'n_valide': n_reponses, 'n_total': n_total,
                    'effectif': k, 'pourcentage': round(100 * k / n_reponses, 1) if n_reponses else np.nan,
                    'ic95_inf': round(lo, 1) if not np.isnan(lo) else np.nan,
                    'ic95_sup': round(hi, 1) if not np.isnan(hi) else np.nan,
                    'statistique': 'proportion (select_multiple)',
                })

        elif xkey in ('period', 'period_raw'):
            vals = pd.to_numeric(serie, errors='coerce').dropna()
            if len(vals) < 3:
                continue
            _, p_norm = stats.shapiro(vals) if len(vals) <= 5000 else (None, 1.0)
            mean, sd = vals.mean(), vals.std()
            sem = sd / np.sqrt(len(vals)) if len(vals) > 1 else 0
            ci_mean = stats.t.interval(1 - ALPHA, len(vals) - 1, loc=mean, scale=sem) if sem > 0 else (mean, mean)
            median = vals.median()
            q1, q3 = vals.quantile([0.25, 0.75])
            rows.append({
                'variable': ind['nom'], 'section': ind['section'], 'label': ind['label'],
                'modalite': 'médiane [IQR] (non normal)' if p_norm < 0.05 else 'moyenne ± ET (normal)',
                'n_valide': len(vals), 'n_total': n_total,
                'effectif': round(mean, 2) if p_norm >= 0.05 else round(median, 2),
                'pourcentage': np.nan,
                'ic95_inf': round(ci_mean[0], 2) if p_norm >= 0.05 else round(q1, 2),
                'ic95_sup': round(ci_mean[1], 2) if p_norm >= 0.05 else round(q3, 2),
                'statistique': f"moyenne={round(mean, 2)} (ET={round(sd, 2)}) | "
                               f"médiane={round(median, 2)} [IQR {round(q1, 2)}-{round(q3, 2)}] | "
                               f"Shapiro p={round(p_norm, 4)}",
            })

    return pd.DataFrame(rows)


# ─────────────────────────────────────────────────────────────────────────
# Scores composites — regroupement piloté par l'utilisateur
# ─────────────────────────────────────────────────────────────────────────

def _item_favorable_series(data, col, valeurs_favorables, is_multi, sens_item):
    serie = data[col]
    favorables = {v.strip() for v in str(valeurs_favorables).split(',') if v.strip()}
    if not favorables:
        return pd.Series(np.nan, index=serie.index)

    if is_multi:
        favorable = serie.apply(
            lambda s: np.nan if pd.isna(s) else bool(set(str(s).split()) & favorables))
    else:
        favorable = serie.apply(lambda v: np.nan if pd.isna(v) else (str(v) in favorables))
    if sens_item == 'negatif':
        favorable = favorable.apply(lambda v: v if pd.isna(v) else (not v))
    return favorable


def compute_composite_scores(dic, data):
    """Calcule un score (%) par domaine défini par l'utilisateur (colonne
    `domaine`), avec dénominateur individualisé (items réellement posés,
    déduits du non-manquant — cohérent avec un skip logic Kobo/ODK)."""
    items = dic[(dic['inclure_score_composite'] == True) & dic['colonne_donnees'].notna()  # noqa: E712
                & (dic['domaine'].astype(str).str.strip() != '')]
    scores = pd.DataFrame(index=data.index)
    detail_par_domaine = {}
    items_ignores = []

    for domaine, grp in items.groupby('domaine'):
        favorables_cols = []
        for _, it in grp.iterrows():
            is_multi = str(it['type_xlsform']).startswith('select_multiple')
            type_simple = str(it['type_xlsform']).split(' ')[0]
            if type_simple not in ('select_one', 'select_multiple'):
                items_ignores.append(it['nom'])
                continue
            fav = _item_favorable_series(data, it['colonne_donnees'], it['valeurs_favorables'],
                                          is_multi, it['sens_item'])
            fav.name = it['nom']
            favorables_cols.append(fav)
        if not favorables_cols:
            continue
        mat = pd.concat(favorables_cols, axis=1)
        denom = mat.notna().sum(axis=1)
        numer = mat.fillna(False).sum(axis=1)
        score = np.where(denom > 0, 100 * numer / denom, np.nan)
        scores[f'Score_{domaine}'] = score
        scores[f'N_items_{domaine}'] = denom
        detail_par_domaine[domaine] = [c.name for c in favorables_cols]

    return scores, detail_par_domaine, items_ignores


def croisement_scores(scores):
    cols = [c for c in scores.columns if c.startswith('Score_')]
    if len(cols) < 2:
        return pd.DataFrame(), pd.DataFrame()
    corr = scores[cols].corr(method='spearman')
    lignes = []
    for i, c1 in enumerate(cols):
        for c2 in cols[i + 1:]:
            sub = scores[[c1, c2]].dropna()
            if len(sub) < 3:
                continue
            rho, p = stats.spearmanr(sub[c1], sub[c2])
            lignes.append({'score_1': c1, 'score_2': c2, 'n': len(sub),
                            'rho_spearman': round(rho, 3), 'p_value': round(p, 4)})
    return corr, pd.DataFrame(lignes)


# ─────────────────────────────────────────────────────────────────────────
# Stratification — pilotée par role == 'stratification' du dictionnaire
# ─────────────────────────────────────────────────────────────────────────

def stratified_analysis(dic, data, scores):
    strat_vars = dic[(dic['role'] == 'stratification') & dic['colonne_donnees'].notna()
                      & dic['type_xlsform'].str.startswith('select_one')]
    score_cols = [c for c in scores.columns if c.startswith('Score_')]
    rows = []

    for _, sv in strat_vars.iterrows():
        strat_label = sv['label'] or sv['nom']
        strat_col = data[sv['colonne_donnees']]
        for score_col in score_cols:
            df = pd.DataFrame({'strat': strat_col, 'score': scores[score_col]}).dropna()
            groupes = [g['score'].values for _, g in df.groupby('strat') if len(g) >= 5]
            noms_groupes = [n for n, g in df.groupby('strat') if len(g) >= 5]
            if len(groupes) < 2:
                continue
            normal = all(len(g) < 5000 and stats.shapiro(g)[1] >= 0.05 for g in groupes if len(g) >= 3)
            if normal:
                stat_val, p = stats.f_oneway(*groupes)
                test_utilise = 'ANOVA'
            else:
                stat_val, p = stats.kruskal(*groupes)
                test_utilise = 'Kruskal-Wallis'
            for nom_g, g in zip(noms_groupes, groupes):
                rows.append({
                    'stratification': strat_label, 'score': score_col, 'groupe': nom_g,
                    'n': len(g), 'moyenne': round(np.mean(g), 1), 'mediane': round(np.median(g), 1),
                    'test': test_utilise, 'p_value_global': round(p, 4),
                    'significatif_5pct': p < 0.05,
                })
    return pd.DataFrame(rows)


# ─────────────────────────────────────────────────────────────────────────
# Qualité des données
# ─────────────────────────────────────────────────────────────────────────

def data_quality_table(dic, data, manquantes, extra_cols):
    rows = []
    reperees = dic[dic['role'].isin(['indicateur', 'stratification'])]

    for _, ind in reperees.iterrows():
        col = ind['colonne_donnees']
        if pd.isna(col):
            rows.append({'variable': ind['nom'], 'relevant_skip_logic': ind.get('relevant_skip_logic', ''),
                          'taux_completion_pct': np.nan, 'n_manquant': np.nan, 'n_aberrant': np.nan,
                          'note': 'ABSENT DES DONNÉES'})
            continue
        serie = data[col]
        n = len(serie)
        n_valid = serie.notna().sum()
        note = ''
        n_aberrant = np.nan

        is_multi = str(ind['type_xlsform']).startswith('select_multiple')
        xkey = _xtype_key(ind['suffixe_phakts'] or None, is_multi)
        type_simple = str(ind['type_xlsform']).split(' ')[0]
        if xkey == 'period' or type_simple in ('integer', 'decimal', 'range'):
            vals = pd.to_numeric(serie, errors='coerce').dropna()
            if len(vals) >= 4:
                q1, q3 = vals.quantile([0.25, 0.75])
                iqr = q3 - q1
                bornes = (q1 - 1.5 * iqr, q3 + 1.5 * iqr)
                n_aberrant = int(((vals < bornes[0]) | (vals > bornes[1])).sum())
                if n_aberrant:
                    note = f"bornes plausibles [{round(bornes[0], 1)}, {round(bornes[1], 1)}] (règle IQR×1.5)"

        rows.append({
            'variable': ind['nom'], 'relevant_skip_logic': ind.get('relevant_skip_logic', ''),
            'taux_completion_pct': round(100 * n_valid / n, 1) if n else np.nan,
            'n_manquant': n - n_valid, 'n_aberrant': n_aberrant, 'note': note,
        })

    df = pd.DataFrame(rows)
    completion_map = dict(zip(df['variable'], df['taux_completion_pct']))
    anomalies_completion = []
    for _, r in df.iterrows():
        if pd.isna(r['taux_completion_pct']) or r['taux_completion_pct'] >= 5:
            continue
        m = re.search(r'\$\{([^}]+)\}', str(r['relevant_skip_logic']))
        if not m:
            continue
        var_condition = m.group(1)
        pct_condition = completion_map.get(var_condition)
        if pct_condition is not None and pct_condition >= 50:
            anomalies_completion.append({
                'variable': r['variable'], 'taux_completion_pct': r['taux_completion_pct'],
                'variable_conditionnante': var_condition, 'taux_completion_conditionnante': pct_condition,
            })
    df = df.drop(columns=['relevant_skip_logic'])

    couverture = {
        'n_variables_dictionnaire': len(reperees),
        'n_non_trouvees_dans_donnees': len(manquantes),
        'variables_non_trouvees': list(manquantes['nom']),
        'n_colonnes_donnees_hors_dictionnaire': len(extra_cols),
        'colonnes_hors_dictionnaire': extra_cols,
        'anomalies_completion': anomalies_completion,
    }
    return df, couverture


# ─────────────────────────────────────────────────────────────────────────
# Figures (Plotly + Kaleido)
# ─────────────────────────────────────────────────────────────────────────

STATUT_COLOR = {
    'zero': '#A0AEC0', 'en_cours': '#E67E22', 'cible': '#2E7D4F',
    'verifier': '#C0392B', 'suivi': '#A0AEC0', 'inconnu': '#D9DEE4',
}


def make_figures(scores, corr, strat_df, groupe_completude=None):
    from modules.report_generator import _fig_to_img
    import plotly.graph_objects as go

    paths = {}

    if groupe_completude:
        gdf = pd.DataFrame(groupe_completude)
        if not gdf.empty:
            colors = [STATUT_COLOR.get(s, '#D9DEE4') for s in gdf['statut']]
            figg = go.Figure(go.Bar(
                x=gdf['groupe'], y=gdf['taux'], marker_color=colors,
                text=[f"{t} %" if pd.notna(t) else '—' for t in gdf['taux']], textposition='outside'))
            figg.update_layout(title="Taux de complétude par groupe", yaxis_title='Taux (%)',
                                plot_bgcolor='white', paper_bgcolor='white',
                                margin=dict(t=40, b=90), xaxis=dict(tickangle=-35))
            p = _fig_to_img(figg.to_json(), width=750, height=380)
            if p:
                paths['completude_groupe'] = p

    cols = [c for c in scores.columns if c.startswith('Score_')]
    if cols:
        fig = go.Figure()
        for c in cols:
            fig.add_trace(go.Box(y=scores[c].dropna(), name=c.replace('Score_', '')))
        fig.update_layout(title="Distribution des scores composites", yaxis_title="Score (%)",
                           plot_bgcolor='white', paper_bgcolor='white', margin=dict(t=40, b=40))
        p = _fig_to_img(fig.to_json(), width=700, height=400)
        if p:
            paths['distribution_scores'] = p

    if not corr.empty:
        fig2 = go.Figure(go.Heatmap(z=corr.values, x=[c.replace('Score_', '') for c in corr.columns],
                                     y=[c.replace('Score_', '') for c in corr.index],
                                     colorscale='RdBu', zmid=0, text=corr.round(2).values,
                                     texttemplate='%{text}'))
        fig2.update_layout(title="Corrélation (Spearman) entre scores composites",
                            plot_bgcolor='white', paper_bgcolor='white', margin=dict(t=40, b=40))
        p2 = _fig_to_img(fig2.to_json(), width=550, height=450)
        if p2:
            paths['correlation_scores'] = p2

    if not strat_df.empty:
        for score_col in strat_df['score'].unique():
            sub = strat_df[strat_df['score'] == score_col]
            for strat_label in sub['stratification'].unique():
                s = sub[sub['stratification'] == strat_label]
                fig3 = go.Figure(go.Bar(x=s['groupe'], y=s['moyenne'],
                                         text=[f"n={n}" for n in s['n']], textposition='outside'))
                fig3.update_layout(
                    title=f"{score_col.replace('Score_', '')} par {strat_label} "
                          f"({s['test'].iloc[0]}, p={s['p_value_global'].iloc[0]})",
                    yaxis_title='Score moyen (%)', plot_bgcolor='white', paper_bgcolor='white',
                    margin=dict(t=60, b=80), xaxis=dict(tickangle=-30))
                key = f"strat_{score_col}_{strat_label}".replace(' ', '_').replace('/', '-')
                p3 = _fig_to_img(fig3.to_json(), width=650, height=380)
                if p3:
                    paths[key] = p3

    return paths


# ─────────────────────────────────────────────────────────────────────────
# Rapport Word
# ─────────────────────────────────────────────────────────────────────────

def build_report(nom_projet, n_soumissions, univ, scores, detail_domaines, items_ignores,
                  corr, corr_tests, strat_df, qual_df, couverture, figures,
                  groupe_completude=None):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from datetime import datetime

    PRIMARY = RGBColor(26, 60, 94)

    doc = Document()
    doc.styles['Normal'].font.size = Pt(10)

    title = doc.add_paragraph(f"ANALYSE ÉPIDÉMIOLOGIQUE — {nom_projet.upper()}")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.bold = True
    title.runs[0].font.size = Pt(16)
    title.runs[0].font.color.rgb = PRIMARY

    sub = doc.add_paragraph(f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')} — "
                             f"{n_soumissions} soumissions")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.italic = True
    sub.runs[0].font.size = Pt(9)
    doc.add_paragraph()

    section = 0

    # Qualité des données / couverture
    section += 1
    h = doc.add_heading(f'{section}. Qualité des données et couverture du dictionnaire', level=1)
    h.runs[0].font.color.rgb = PRIMARY
    if couverture['variables_non_trouvees']:
        doc.add_paragraph(
            f"{couverture['n_non_trouvees_dans_donnees']} variable(s) du dictionnaire non retrouvée(s) "
            f"dans les données réellement collectées :"
        )
        for v in couverture['variables_non_trouvees']:
            doc.add_paragraph(v, style='List Bullet')
    if couverture['colonnes_hors_dictionnaire']:
        doc.add_paragraph(f"{couverture['n_colonnes_donnees_hors_dictionnaire']} colonne(s) présente(s) "
                           f"dans les données mais absente(s) du dictionnaire :")
        for c in couverture['colonnes_hors_dictionnaire']:
            doc.add_paragraph(c, style='List Bullet')
    n_aberrants = qual_df['n_aberrant'].fillna(0).sum() if not qual_df.empty else 0
    moy_completion = round(qual_df['taux_completion_pct'].mean(skipna=True), 1) if not qual_df.empty else None
    doc.add_paragraph(
        f"Complétion moyenne des indicateurs disponibles : {moy_completion if moy_completion is not None else '—'} %. "
        f"{int(n_aberrants)} valeur(s) aberrante(s) détectée(s) au total (règle IQR×1.5)."
    )
    anomalies_c = couverture.get('anomalies_completion', [])
    if anomalies_c:
        alerte = doc.add_paragraph()
        r = alerte.add_run(
            f"⚠ ALERTE — {len(anomalies_c)} indicateur(s) quasi vide(s) (<5 % de complétion) alors "
            f"que la variable qui les conditionne est bien renseignée :"
        )
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        for a in anomalies_c:
            doc.add_paragraph(
                f"{a['variable']} : {a['taux_completion_pct']} % de complétion, alors que "
                f"{a['variable_conditionnante']} est rempli à {a['taux_completion_conditionnante']} % "
                f"— vérifier la logique de saut réellement déployée.", style='List Bullet')

    # Complétude par groupe (réutilise modules/projets.py — pas recalculée ici)
    if groupe_completude:
        section += 1
        doc.add_page_break()
        h = doc.add_heading(f'{section}. Taux de complétude par groupe', level=1)
        h.runs[0].font.color.rgb = PRIMARY
        if 'completude_groupe' in figures:
            doc.add_picture(figures['completude_groupe'], width=Inches(6.2))
        gdf = pd.DataFrame(groupe_completude)
        if not gdf.empty:
            gtable = doc.add_table(rows=1, cols=4)
            gtable.style = 'Table Grid'
            for i, hdr in enumerate(['Groupe', 'Reçu', 'Cible', 'Taux (%)']):
                gtable.rows[0].cells[i].text = hdr
            for _, r in gdf.iterrows():
                cells = gtable.add_row().cells
                cells[0].text = str(r['groupe'])
                cells[1].text = str(r['recu'])
                cells[2].text = str(r['cible'])
                cells[3].text = '' if pd.isna(r['taux']) else str(r['taux'])
                for c in cells:
                    for p in c.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(8)

    # Scores composites
    if detail_domaines:
        section += 1
        doc.add_page_break()
        h = doc.add_heading(f'{section}. Scores composites', level=1)
        h.runs[0].font.color.rgb = PRIMARY
        for domaine, items in detail_domaines.items():
            col = f'Score_{domaine}'
            if col not in scores.columns:
                continue
            vals = scores[col].dropna()
            if vals.empty:
                continue
            doc.add_heading(f"{domaine} ({len(items)} item(s))", level=2)
            doc.add_paragraph(f"N = {len(vals)} · Moyenne = {round(vals.mean(), 1)} % · "
                               f"Médiane = {round(vals.median(), 1)} % · "
                               f"IC95% de la moyenne = [{round(vals.mean() - Z*vals.sem(), 1)}, "
                               f"{round(vals.mean() + Z*vals.sem(), 1)}]")
            doc.add_paragraph("Items : " + ", ".join(items))
        if items_ignores:
            doc.add_paragraph(
                "Items ignorés dans les scores (type non catégoriel, non pris en charge pour un "
                "score composite) : " + ", ".join(items_ignores))
        if 'distribution_scores' in figures:
            doc.add_picture(figures['distribution_scores'], width=Inches(6))

        # Croisement
        if not corr_tests.empty:
            section += 1
            h = doc.add_heading(f'{section}. Croisement des scores composites', level=1)
            h.runs[0].font.color.rgb = PRIMARY
            if 'correlation_scores' in figures:
                doc.add_picture(figures['correlation_scores'], width=Inches(5))
            for _, r in corr_tests.iterrows():
                sig = ' (significatif à 5 %)' if r['p_value'] < 0.05 else ''
                doc.add_paragraph(
                    f"{r['score_1'].replace('Score_', '')} × {r['score_2'].replace('Score_', '')} : "
                    f"rho={r['rho_spearman']}, p={r['p_value']}{sig} (n={r['n']})", style='List Bullet')

        # Stratification
        if not strat_df.empty:
            section += 1
            doc.add_page_break()
            h = doc.add_heading(f'{section}. Stratification des scores composites', level=1)
            h.runs[0].font.color.rgb = PRIMARY
            for strat_label in strat_df['stratification'].unique():
                doc.add_heading(strat_label, level=2)
                sub = strat_df[strat_df['stratification'] == strat_label]
                for score_col in sub['score'].unique():
                    s = sub[sub['score'] == score_col]
                    p = s['p_value_global'].iloc[0]
                    sig = ' — différence significative' if p < 0.05 else ' — pas de différence significative'
                    doc.add_paragraph(f"{score_col.replace('Score_', '')} ({s['test'].iloc[0]}, p={p}){sig}")
                    key = f"strat_{score_col}_{strat_label}".replace(' ', '_').replace('/', '-')
                    if key in figures:
                        doc.add_picture(figures[key], width=Inches(5.5))

    # Statistiques univariées (extrait)
    section += 1
    doc.add_page_break()
    h = doc.add_heading(f'{section}. Statistiques univariées par indicateur (extrait — CSV complet fourni)', level=1)
    h.runs[0].font.color.rgb = PRIMARY
    if not univ.empty:
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        for i, hd in enumerate(['Variable', 'Section', 'Modalité', 'N', '%', 'IC95%']):
            table.rows[0].cells[i].text = hd
        for _, r in univ.head(60).iterrows():
            cells = table.add_row().cells
            cells[0].text = str(r['variable'])
            cells[1].text = str(r['section'])[:30]
            cells[2].text = str(r['modalite'])
            cells[3].text = str(r['n_valide'])
            cells[4].text = '' if pd.isna(r['pourcentage']) else str(r['pourcentage'])
            cells[5].text = '' if pd.isna(r['ic95_inf']) else f"[{r['ic95_inf']}, {r['ic95_sup']}]"
            for c in cells:
                for p in c.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(8)
    else:
        doc.add_paragraph("Aucun indicateur univarié calculable (voir couverture dictionnaire ↔ données).")

    footer = doc.add_paragraph(
        "Rapport généré automatiquement par SPAD PHAKTS Analyzer. Les seuils de significativité (5 %) "
        "et les choix de test (paramétrique/non paramétrique) sont déterminés automatiquement par test "
        "de normalité (Shapiro-Wilk) ; à revoir par un⋅e biostatisticien⋅ne pour toute publication."
    )
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
