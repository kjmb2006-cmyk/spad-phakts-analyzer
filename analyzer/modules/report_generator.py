import os
import re
import io
from datetime import datetime
from io import StringIO
import pandas as pd
import numpy as np
import plotly.io as pio
import tempfile

# Contournement bug Kaleido : son script shell fait `cd $DIR` sans guillemets,
# ce qui plante quand le chemin contient des espaces ("Consultance 2026").
# Stratégie : on cherche plotly.min.js dans cet ordre :
#   1. /tmp/plotly_spad.min.js  — copie faite au démarrage (hors-ligne ✓)
#   2. static/vendor/plotly.min.js — bundlé dans le projet (hors-ligne ✓)
#   3. CDN Plotly                  — fallback si Internet disponible
def _setup_kaleido():
    import shutil, pathlib
    tmp_js = pathlib.Path('/tmp/plotly_spad.min.js')
    # Copie depuis static/vendor si absent de /tmp
    vendor = pathlib.Path(__file__).parent.parent / 'static' / 'vendor' / 'plotly.min.js'
    if not tmp_js.exists() and vendor.exists():
        shutil.copy2(str(vendor), str(tmp_js))
    if tmp_js.exists():
        try:
            pio.kaleido.scope.plotlyjs = str(tmp_js)
            return
        except Exception:
            pass
    # Dernier recours : CDN (nécessite Internet)
    try:
        pio.kaleido.scope.plotlyjs = 'https://cdn.plot.ly/plotly-latest.min.js'
    except Exception:
        pass

_setup_kaleido()

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                  TableStyle, PageBreak, HRFlowable, Image)
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.pdfgen import canvas

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from modules.data_loader import get_var_types
from modules.descriptive import tris_a_plat, statistiques_continues, analyse_binaire_groupe
from modules.crosstabs import tableau_croise
from modules.raw_analysis import (
    overview_stats, systematic_analysis, data_quality_gauge,
    missing_heatmap, missing_bar_chart, distributions_chart, composition_chart,
    descriptive_summary, correlation_matrix,
)
from modules.comments import (auto_comment_categorical, auto_comment_continuous,
                              auto_comment_crosstab, auto_comment_binary_group)

PRIMARY = colors.HexColor('#1a3c5e')
ACCENT  = colors.HexColor('#e05c1a')
LIGHT   = colors.HexColor('#f0f6fc')
GRAY    = colors.HexColor('#6c757d')
WHITE   = colors.white


def _styles():
    return {
        'Title':     ParagraphStyle('Title', fontSize=22, textColor=PRIMARY,
                                    alignment=TA_CENTER, spaceAfter=6,
                                    fontName='Helvetica-Bold'),
        'Subtitle':  ParagraphStyle('Subtitle', fontSize=12, textColor=GRAY,
                                    alignment=TA_CENTER, spaceAfter=12),
        'H1':        ParagraphStyle('H1', fontSize=15, textColor=PRIMARY,
                                    spaceBefore=14, spaceAfter=6,
                                    fontName='Helvetica-Bold'),
        'H2':        ParagraphStyle('H2', fontSize=12, textColor=ACCENT,
                                    spaceBefore=10, spaceAfter=4,
                                    fontName='Helvetica-Bold'),
        'Body':      ParagraphStyle('Body', fontSize=10, spaceAfter=4, leading=14),
        'Meta':      ParagraphStyle('Meta', fontSize=9, textColor=GRAY, spaceAfter=2),
        'Caption':   ParagraphStyle('Caption', fontSize=8, textColor=GRAY,
                                    alignment=TA_CENTER, spaceAfter=8),
        'Highlight': ParagraphStyle('Highlight', fontSize=10, textColor=PRIMARY,
                                    backColor=LIGHT, borderPad=6, spaceAfter=6),
    }


def _df_to_table(df: pd.DataFrame, max_rows: int = 40) -> Table:
    if len(df) > max_rows:
        df = df.head(max_rows)
    data = [[str(c) for c in row]
            for row in [df.columns.tolist()] + df.values.tolist()]
    t = Table(data, repeatRows=1)
    t.setStyle(TableStyle([
        ('BACKGROUND',    (0, 0), (-1, 0),  PRIMARY),
        ('TEXTCOLOR',     (0, 0), (-1, 0),  WHITE),
        ('FONTNAME',      (0, 0), (-1, 0),  'Helvetica-Bold'),
        ('FONTSIZE',      (0, 0), (-1, -1), 8),
        ('ALIGN',         (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
        ('ROWBACKGROUNDS',(0, 1), (-1, -1), [WHITE, LIGHT]),
        ('GRID',          (0, 0), (-1, -1), 0.4, colors.HexColor('#dee2e6')),
        ('TOPPADDING',    (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('LEFTPADDING',   (0, 0), (-1, -1), 6),
    ]))
    return t


def _fig_to_img(fig_json: str, width=520, height=300) -> str | None:
    try:
        fig = pio.from_json(fig_json)
        img = fig.to_image(format='png', width=width, height=height, scale=1.5)
        tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
        tmp.write(img)
        tmp.close()
        return tmp.name
    except Exception:
        return None


class _NumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        n = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self._footer(n)
            canvas.Canvas.showPage(self)
        canvas.Canvas.save(self)

    def _footer(self, total):
        self.saveState()
        self.setFont('Helvetica', 8)
        self.setFillColor(GRAY)
        self.drawString(2*cm, 1.2*cm, 'SPAD Analyzer — Rapport automatique')
        self.drawRightString(A4[0]-2*cm, 1.2*cm, f'Page {self._pageNumber} / {total}')
        self.setStrokeColor(colors.HexColor('#dee2e6'))
        self.line(2*cm, 1.5*cm, A4[0]-2*cm, 1.5*cm)
        self.restoreState()


def generate_pdf_report(df: pd.DataFrame, output_path: str = None,
                         title: str = 'Rapport', author: str = 'Analyste',
                         selected_analyses: list = None,
                         cat_vars: list = None, num_vars: list = None,
                         bin_groups: list = None,
                         crosstabs_row=None,
                         crosstabs_col=None,
                         user_comments: dict = None,
                         multi_survey_results: list = None,
                         multi_survey_meta: dict = None):
    S = _styles()
    tmp_imgs = []
    bin_groups = bin_groups or []
    user_comments = user_comments or {}

    # Style dédié pour les commentaires
    if 'CommentAuto' not in S:
        S['CommentAuto'] = ParagraphStyle(
            'CommentAuto', fontSize=9, textColor=PRIMARY,
            backColor=LIGHT, borderPad=6, spaceAfter=4, leading=12,
            leftIndent=8, rightIndent=8)
        S['CommentUser'] = ParagraphStyle(
            'CommentUser', fontSize=9.5, textColor=colors.HexColor('#2C3E50'),
            spaceAfter=8, leading=13, leftIndent=8, rightIndent=8,
            fontName='Helvetica-Oblique')

    def _add_comment(story, auto: str, user: str):
        if auto:
            story.append(Paragraph(f'<b>Lecture automatique :</b> {auto}', S['CommentAuto']))
        if user:
            story.append(Paragraph(f'<b>Commentaire de l\'analyste :</b> {user}', S['CommentUser']))
    # Activer l'analyse brute par défaut si l'utilisateur n'a rien précisé
    if selected_analyses is None:
        selected_analyses = ['raw']
    cat_vars = cat_vars or []
    num_vars = num_vars or []

    buf = io.BytesIO() if output_path is None else None
    dest = buf if buf is not None else output_path

    doc = SimpleDocTemplate(dest, pagesize=A4,
                             topMargin=2.5*cm, bottomMargin=2.2*cm,
                             leftMargin=2*cm, rightMargin=2*cm)
    story = []
    var_types = get_var_types(df)

    # ── Page de garde ────────────────────────────────────────────────────────
    story += [
        Spacer(1, 3*cm),
        Paragraph(title, S['Title']),
        Spacer(1, 0.3*cm),
        Paragraph('Rapport d\'analyse statistique — SPAD Analyzer', S['Subtitle']),
        HRFlowable(width='90%', thickness=2, color=PRIMARY, spaceAfter=12),
        Spacer(1, 0.4*cm),
    ]

    meta_rows = [
        ['Auteur', author],
        ['Date', datetime.now().strftime('%d/%m/%Y à %H:%M')],
        ['Observations', str(df.shape[0])],
        ['Variables analysées', str(df.shape[1])],
        ['Données manquantes',
         f"{int(df.isnull().sum().sum())} cellules"],
    ]
    mt = Table([[Paragraph(k, S['Body']), Paragraph(v, S['Body'])]
                for k, v in meta_rows], colWidths=[6*cm, 9*cm])
    mt.setStyle(TableStyle([
        ('BACKGROUND',  (0, 0), (0, -1), LIGHT),
        ('FONTNAME',    (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE',    (0, 0), (-1,-1), 10),
        ('GRID',        (0, 0), (-1,-1), 0.5, colors.HexColor('#dee2e6')),
        ('TOPPADDING',  (0, 0), (-1,-1), 5),
        ('BOTTOMPADDING',(0,0), (-1,-1), 5),
        ('LEFTPADDING', (0, 0), (-1,-1), 8),
    ]))
    story += [mt, PageBreak()]

    section = 1

    # ── 1. Questions à choix multiples ───────────────────────────────────────
    if 'binaire' in selected_analyses and bin_groups:
        story += [
            Paragraph(f'{section}. Questions à choix multiples', S['H1']),
            HRFlowable(width='100%', thickness=1, color=ACCENT, spaceAfter=8),
        ]
        for g in bin_groups:
            story.append(Paragraph(g['parent'], S['H2']))
            try:
                r = analyse_binaire_groupe(df, g['vars'])
                tbl_df = pd.read_html(StringIO(r['table_html']))[0]
                story.append(_df_to_table(tbl_df))
                story.append(Spacer(1, 0.3*cm))
                img = _fig_to_img(r['chart'], width=520, height=max(250, 35*len(g['vars'])+80))
                if img:
                    tmp_imgs.append(img)
                    story.append(Image(img, width=14*cm,
                                       height=min(10*cm, max(6*cm, 0.9*len(g['vars'])*cm))))
                    story.append(Paragraph(f'Figure : {g["parent"]}', S['Caption']))
                _add_comment(story,
                             auto_comment_binary_group({'parent': g['parent'], **r}),
                             user_comments.get(f'descriptive_bin:{g["parent"]}', ''))
            except Exception as e:
                story.append(Paragraph(f'Erreur : {e}', S['Body']))
            story.append(Spacer(1, 0.5*cm))
        section += 1
        story.append(PageBreak())

    # ── 2. Statistiques descriptives catégorielles ───────────────────────────
    if 'descriptive' in selected_analyses and cat_vars:
        story += [
            Paragraph(f'{section}. Statistiques descriptives — Variables catégorielles', S['H1']),
            HRFlowable(width='100%', thickness=1, color=ACCENT, spaceAfter=8),
        ]
        for var in cat_vars:
            if var not in df.columns:
                continue
            story.append(Paragraph(f'Tri à plat — {var}', S['H2']))
            try:
                r = tris_a_plat(df, var)
                story.append(Paragraph(
                    f'N valides = <b>{r["n_valid"]}</b> | Manquantes : <b>{r["missing"]}</b> ({r["missing_pct"]}%)',
                    S['Meta']))
                tbl_df = pd.read_html(StringIO(r['table_html']))[0]
                story.append(_df_to_table(tbl_df))
                story.append(Spacer(1, 0.3*cm))
                img = _fig_to_img(r['chart_bar'], width=520, height=280)
                if img:
                    tmp_imgs.append(img)
                    story.append(Image(img, width=14*cm, height=7*cm))
                    story.append(Paragraph(f'Figure : Distribution de {var}', S['Caption']))
                _add_comment(story,
                             auto_comment_categorical(r),
                             user_comments.get(f'descriptive_cat:{var}', ''))
            except Exception as e:
                story.append(Paragraph(f'Erreur : {e}', S['Body']))
            story.append(Spacer(1, 0.5*cm))
        section += 1
        story.append(PageBreak())

    # ── 3. Statistiques descriptives continues ───────────────────────────────
    if 'descriptive' in selected_analyses and num_vars:
        story += [
            Paragraph(f'{section}. Statistiques descriptives — Variables continues', S['H1']),
            HRFlowable(width='100%', thickness=1, color=ACCENT, spaceAfter=8),
        ]
        for var in num_vars:
            if var not in df.columns:
                continue
            story.append(Paragraph(f'Distribution — {var}', S['H2']))
            try:
                r = statistiques_continues(df, var)
                tbl_df = pd.read_html(StringIO(r['table_html']))[0]
                story.append(_df_to_table(tbl_df))
                story.append(Spacer(1, 0.3*cm))
                img = _fig_to_img(r['chart_hist'], width=520, height=300)
                if img:
                    tmp_imgs.append(img)
                    story.append(Image(img, width=14*cm, height=7.5*cm))
                    story.append(Paragraph(f'Figure : Distribution de {var}', S['Caption']))
                _add_comment(story,
                             auto_comment_continuous(r),
                             user_comments.get(f'descriptive_num:{var}', ''))
            except Exception as e:
                story.append(Paragraph(f'Erreur : {e}', S['Body']))
            story.append(Spacer(1, 0.5*cm))
        section += 1
        story.append(PageBreak())

    # ── 4. Tableaux croisés ──────────────────────────────────────────────────
    all_cat = cat_vars or [c for c, t in var_types.items() if t == 'categorielle']
    if 'crosstabs' in selected_analyses and len(all_cat) >= 2:
        story += [
            Paragraph(f'{section}. Tableaux croisés', S['H1']),
            HRFlowable(width='100%', thickness=1, color=ACCENT, spaceAfter=8),
        ]
        row_var = crosstabs_row if crosstabs_row in all_cat else all_cat[0]
        col_var = crosstabs_col if crosstabs_col in all_cat else (all_cat[1] if len(all_cat) > 1 else all_cat[0])
        if row_var == col_var and len(all_cat) > 1:
            col_var = all_cat[1] if all_cat[0] == row_var else all_cat[0]

        try:
            r = tableau_croise(df, row_var, col_var, 'row')
            story.append(Paragraph(
                f'<b>Variables</b> : {row_var} × {col_var}', S['Meta']))
            story.append(Paragraph(
                f'<b>χ²</b> = {r["chi2"]} | <b>p-value</b> = {r["p_val"]} | '
                f'<b>V de Cramér</b> = {r["cramers_v"]} — '
                f'Liaison <b>{r["assoc_label"]}</b> ({r["sig_label"]})',
                S['Highlight']))
            ct = pd.crosstab(df[row_var].dropna().astype(str),
                              df[col_var].dropna().astype(str))
            story.append(_df_to_table(ct.reset_index()))
            img = _fig_to_img(r['chart_bar'], width=520, height=320)
            if img:
                tmp_imgs.append(img)
                story.append(Image(img, width=14*cm, height=8*cm))
            _add_comment(story,
                         auto_comment_crosstab(r),
                         user_comments.get(f'crosstabs:{row_var}|{col_var}', ''))
        except Exception as e:
            story.append(Paragraph(f'Erreur : {e}', S['Body']))
        section += 1
        story.append(PageBreak())

    # ── 5. Analyse brute (overview & data quality) ───────────────────────────
    if 'raw' in selected_analyses:
        story += [
            Paragraph(f'{section}. Analyse brute', S['H1']),
            HRFlowable(width='100%', thickness=1, color=ACCENT, spaceAfter=8),
        ]
        try:
            sysa  = systematic_analysis(df)
            desc  = descriptive_summary(df)
            ov    = overview_stats(df)

            # ── 5.1 Qualité & Profil ─────────────────────────────────────────
            story.append(Paragraph(f'{section}.1 Qualité & Profil', S['H2']))
            q_img = _fig_to_img(data_quality_gauge(sysa.get('quality', {})), width=420, height=240)
            if q_img:
                tmp_imgs.append(q_img)
                story.append(Image(q_img, width=11*cm, height=5.5*cm))
                story.append(Paragraph('Figure : Score de qualité des données', S['Caption']))
            comp_img = _fig_to_img(composition_chart(sysa), width=420, height=240)
            if comp_img:
                tmp_imgs.append(comp_img)
                story.append(Image(comp_img, width=11*cm, height=5.5*cm))
                story.append(Paragraph('Figure : Composition des variables', S['Caption']))
            prof = sysa.get('profile', {}) or {}
            if prof:
                prof_rows = [
                    ['Observations', prof.get('n_obs', '—')],
                    ['Variables', prof.get('n_vars', '—')],
                    ['Complètes', prof.get('complete_vars', '—')],
                    ['Avec manquantes', prof.get('vars_with_missing', '—')],
                    ['Continues', prof.get('n_continuous', '—')],
                    ['Catégorielles', prof.get('n_categorical', '—')],
                    ['Binaires', prof.get('n_binary', '—')],
                ]
                story.append(_df_to_table(pd.DataFrame(prof_rows, columns=['Indicateur', 'Valeur'])))
                story.append(Spacer(1, 0.3*cm))

            # ── 5.2 Variables continues ──────────────────────────────────────
            story.append(PageBreak())
            story.append(Paragraph(f'{section}.2 Variables continues', S['H2']))
            cont_list = desc.get('continuous', []) or []
            if not cont_list:
                story.append(Paragraph('Aucune variable continue détectée.', S['Body']))
            for r in cont_list[:10]:
                story.append(Paragraph(f"<b>{r['var']}</b>", S['Body']))
                stats_df = pd.DataFrame(list(r['stats'].items()), columns=['Statistique', 'Valeur'])
                story.append(_df_to_table(stats_df))
                story.append(Spacer(1, 0.2*cm))
                h_img = _fig_to_img(r.get('chart_hist'), width=480, height=260)
                if h_img:
                    tmp_imgs.append(h_img)
                    story.append(Image(h_img, width=14*cm, height=6.5*cm))
                b_img = _fig_to_img(r.get('chart_box'), width=480, height=260)
                if b_img:
                    tmp_imgs.append(b_img)
                    story.append(Image(b_img, width=14*cm, height=6.5*cm))
                story.append(Spacer(1, 0.4*cm))

            # ── 5.3 Catégorielles ────────────────────────────────────────────
            story.append(PageBreak())
            story.append(Paragraph(f'{section}.3 Catégorielles', S['H2']))
            cat_list = desc.get('categorical', []) or []
            if not cat_list:
                story.append(Paragraph('Aucune variable catégorielle détectée.', S['Body']))
            for r in cat_list[:10]:
                story.append(Paragraph(f"<b>{r['variable']}</b> — {r['n_categories']} catégorie(s), "
                                        f"dominante : <b>{r['dominant']}</b> ({r['dominant_pct']}%)", S['Body']))
                try:
                    freq_df = pd.read_html(StringIO(r['table']))[0]
                    story.append(_df_to_table(freq_df))
                except Exception:
                    pass
                story.append(Spacer(1, 0.2*cm))
                bar_img = _fig_to_img(r.get('chart_bar'), width=480, height=300)
                if bar_img:
                    tmp_imgs.append(bar_img)
                    story.append(Image(bar_img, width=14*cm, height=7*cm))
                story.append(Spacer(1, 0.4*cm))

            # ── 5.4 Binaires ─────────────────────────────────────────────────
            story.append(PageBreak())
            story.append(Paragraph(f'{section}.4 Binaires', S['H2']))
            bin_list = desc.get('binary', []) or []
            if not bin_list:
                story.append(Paragraph('Aucune variable binaire détectée.', S['Body']))
            else:
                # Tableau récapitulatif
                bin_rows = [[r['variable'], r['n_valid'], r['n_missing'],
                             f"{r['pct_missing']}%", r['label_pos'],
                             f"{r['prevalence']}%"] for r in bin_list]
                bin_df = pd.DataFrame(bin_rows, columns=[
                    'Variable', 'N valides', 'Manquantes', '% manq.',
                    'Modalité positive', 'Prévalence',
                ])
                story.append(_df_to_table(bin_df))
                story.append(Spacer(1, 0.3*cm))

            # ── 5.5 Manquantes ───────────────────────────────────────────────
            story.append(PageBreak())
            story.append(Paragraph(f'{section}.5 Manquantes', S['H2']))
            try:
                miss_bar = missing_bar_chart(df)
                if miss_bar:
                    bar_img = _fig_to_img(miss_bar, width=520, height=300)
                    if bar_img:
                        tmp_imgs.append(bar_img)
                        story.append(Image(bar_img, width=15*cm, height=7*cm))
                        story.append(Paragraph('Figure : Taux de données manquantes par variable', S['Caption']))
            except Exception:
                pass
            try:
                miss_heat = missing_heatmap(df)
                if miss_heat:
                    h_img = _fig_to_img(miss_heat, width=520, height=300)
                    if h_img:
                        tmp_imgs.append(h_img)
                        story.append(Image(h_img, width=15*cm, height=7*cm))
                        story.append(Paragraph('Figure : Carte des données manquantes', S['Caption']))
            except Exception:
                pass

            # ── 5.6 Corrélations ─────────────────────────────────────────────
            story.append(PageBreak())
            story.append(Paragraph(f'{section}.6 Corrélations', S['H2']))
            try:
                corr_fig = correlation_matrix(df)
                if corr_fig:
                    c_img = _fig_to_img(corr_fig, width=520, height=400)
                    if c_img:
                        tmp_imgs.append(c_img)
                        story.append(Image(c_img, width=14*cm, height=10*cm))
                        story.append(Paragraph('Figure : Matrice de corrélations (variables numériques)', S['Caption']))
                    else:
                        story.append(Paragraph('Pas assez de variables numériques pour une matrice de corrélations.', S['Body']))
                else:
                    story.append(Paragraph('Pas assez de variables numériques pour une matrice de corrélations.', S['Body']))
            except Exception:
                story.append(Paragraph('Pas assez de variables numériques pour une matrice de corrélations.', S['Body']))

            # ── 5.7 Vue d'ensemble ───────────────────────────────────────────
            story.append(PageBreak())
            story.append(Paragraph(f'{section}.7 Vue d\'ensemble', S['H2']))
            try:
                tbl_df = pd.read_html(StringIO(ov['table_html']))[0]
                story.append(_df_to_table(tbl_df, max_rows=120))
            except Exception as e:
                story.append(Paragraph(f'Vue d\'ensemble indisponible : {e}', S['Body']))

            # ── Annexe : anomalies & recommandations (si présentes) ──────────
            if sysa.get('anomalies') or sysa.get('recommendations'):
                story.append(Spacer(1, 0.4*cm))
                story.append(Paragraph(f'{section}.8 Anomalies & Recommandations', S['H2']))
                if sysa.get('anomalies'):
                    for a in sysa['anomalies'][:20]:
                        story.append(Paragraph(f"- {a['type']} — {a['variable']} : {a['note']}", S['Body']))
                if sysa.get('recommendations'):
                    story.append(Spacer(1, 0.2*cm))
                    story.append(Paragraph('<b>Recommandations</b>', S['Body']))
                    for r in sysa['recommendations'][:10]:
                        story.append(Paragraph(f"• {r['type']} — {r['rationale']}", S['Body']))
        except Exception as e:
            story.append(Paragraph(f'Erreur lors de l\'analyse brute : {e}', S['Body']))
        section += 1
        story.append(PageBreak())

    # ── 6. Multi-enquête (DPF / PHAKTS) ──────────────────────────────────────
    if 'multi_survey' in selected_analyses and multi_survey_results:
        story += [
            Paragraph(f'{section}. Analyse multi-enquête (DPF / PHAKTS)', S['H1']),
            HRFlowable(width='100%', thickness=1, color=ACCENT, spaceAfter=8),
        ]
        if multi_survey_meta:
            story.append(Paragraph(
                f"<b>{multi_survey_meta['n_surveys']}</b> enquête(s) fusionnée(s) "
                f"selon le mode <b>{multi_survey_meta['mode']}</b> — "
                f"<b>{multi_survey_meta['n_obs']}</b> observations, "
                f"<b>{multi_survey_meta['n_vars']}</b> variables alignées.", S['Body']))
            story.append(Spacer(1, 0.3*cm))

        for idx, r in enumerate(multi_survey_results, start=1):
            story.append(Paragraph(
                f"{section}.{idx} Variable <b>{r.get('variable', '?')}</b>", S['H2']))
            if r.get('error'):
                story.append(Paragraph(f"Erreur : {r['error']}", S['Body']))
                story.append(Spacer(1, 0.3*cm))
                continue
            try:
                if r.get('kind') == 'continuous' and r.get('summary_html'):
                    story.append(Paragraph('Résumé statistique par enquête :', S['Body']))
                    summ_df = pd.read_html(StringIO(r['summary_html']))[0]
                    story.append(_df_to_table(summ_df))
                elif r.get('kind') == 'categorical':
                    if r.get('count_html'):
                        story.append(Paragraph('Effectifs :', S['Body']))
                        c_df = pd.read_html(StringIO(r['count_html']))[0]
                        story.append(_df_to_table(c_df))
                    if r.get('pct_html'):
                        story.append(Spacer(1, 0.2*cm))
                        story.append(Paragraph('Pourcentages (% colonne) :', S['Body']))
                        p_df = pd.read_html(StringIO(r['pct_html']))[0]
                        story.append(_df_to_table(p_df))
            except Exception as e:
                story.append(Paragraph(f'Tableaux indisponibles : {e}', S['Body']))
            # Graphique
            chart_img = _fig_to_img(r.get('chart'), width=520, height=320)
            if chart_img:
                tmp_imgs.append(chart_img)
                story.append(Spacer(1, 0.2*cm))
                story.append(Image(chart_img, width=15*cm, height=8*cm))
                story.append(Paragraph(
                    f"Figure : Comparaison de <b>{r.get('variable', '?')}</b> entre enquêtes",
                    S['Caption']))
            story.append(Spacer(1, 0.4*cm))
        section += 1
        story.append(PageBreak())

    # ── Conclusion ───────────────────────────────────────────────────────────
    story += [
        Paragraph('Conclusion', S['H1']),
        HRFlowable(width='100%', thickness=1, color=ACCENT, spaceAfter=8),
        Paragraph(
            f'Ce rapport a été généré automatiquement le '
            f'{datetime.now().strftime("%d/%m/%Y")} par <b>SPAD Analyzer</b>. '
            f'Il couvre <b>{df.shape[0]}</b> observations et '
            f'<b>{df.shape[1]}</b> variables.', S['Body']),
    ]

    doc.build(story, canvasmaker=_NumberedCanvas)

    for p in tmp_imgs:
        try:
            os.remove(p)
        except Exception:
            pass

    if buf is not None:
        buf.seek(0)
        return buf
    return None


def generate_word_report(df: pd.DataFrame, output_path: str = None,
                          title: str = 'Rapport', author: str = 'Analyste',
                          selected_analyses: list = None,
                          cat_vars: list = None, num_vars: list = None,
                          bin_groups: list = None,
                          crosstabs_row=None,
                          crosstabs_col=None,
                          user_comments: dict = None,
                          multi_survey_results: list = None,
                          multi_survey_meta: dict = None):
    tmp_imgs = []
    bin_groups = bin_groups or []
    user_comments = user_comments or {}

    def _add_comment_word(doc, auto: str, user: str):
        if auto:
            p = doc.add_paragraph()
            r = p.add_run('Lecture automatique : ')
            r.bold = True
            p.add_run(auto)
        if user:
            p = doc.add_paragraph()
            r = p.add_run('Commentaire de l\'analyste : ')
            r.bold = True
            r2 = p.add_run(user)
            r2.italic = True
    # Activer l'analyse brute par défaut si l'utilisateur n'a rien précisé
    if selected_analyses is None:
        selected_analyses = ['raw']
    cat_vars = cat_vars or []
    num_vars = num_vars or []

    doc = Document()

    # Styles
    title_style = doc.styles['Title']
    title_style.font.size = Pt(22)
    title_style.font.color.rgb = RGBColor(26, 60, 94)

    heading1_style = doc.styles['Heading 1']
    heading1_style.font.size = Pt(15)
    heading1_style.font.color.rgb = RGBColor(26, 60, 94)

    heading2_style = doc.styles['Heading 2']
    heading2_style.font.size = Pt(12)
    heading2_style.font.color.rgb = RGBColor(224, 92, 26)

    normal_style = doc.styles['Normal']
    normal_style.font.size = Pt(10)

    var_types = get_var_types(df)

    # ── Page de garde ────────────────────────────────────────────────────────
    title_para = doc.add_paragraph(title, style='Title')
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Rapport d\'analyse statistique — SPAD Analyzer')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style = doc.styles['Subtitle'] if 'Subtitle' in doc.styles else normal_style

    doc.add_paragraph()  # Spacer

    # Metadata table
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.style = 'Table Grid'
    meta_rows = [
        ['Auteur', author],
        ['Date', datetime.now().strftime('%d/%m/%Y à %H:%M')],
        ['Observations', str(df.shape[0])],
        ['Variables analysées', str(df.shape[1])],
        ['Données manquantes', f"{int(df.isnull().sum().sum())} cellules"],
    ]
    for i, (k, v) in enumerate(meta_rows):
        cells = meta_table.rows[i].cells
        cells[0].text = k
        cells[1].text = v
        cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_page_break()
    section = 1

    # ── 1. Questions à choix multiples ───────────────────────────────────────
    if 'binaire' in selected_analyses and bin_groups:
        doc.add_heading(f'{section}. Questions à choix multiples', level=1)
        for g in bin_groups:
            doc.add_heading(g['parent'], level=2)
            try:
                r = analyse_binaire_groupe(df, g['vars'])
                tbl_df = pd.read_html(StringIO(r['table_html']))[0]

                # Add table
                table = doc.add_table(rows=len(tbl_df)+1, cols=len(tbl_df.columns))
                table.style = 'Table Grid'

                # Header
                hdr_cells = table.rows[0].cells
                for j, col in enumerate(tbl_df.columns):
                    hdr_cells[j].text = str(col)
                    hdr_cells[j].paragraphs[0].runs[0].font.bold = True

                # Data
                for i, row in enumerate(tbl_df.values):
                    row_cells = table.rows[i+1].cells
                    for j, val in enumerate(row):
                        row_cells[j].text = str(val)

                doc.add_paragraph()

                # Add chart if available
                img = _fig_to_img(r['chart'], width=520, height=max(250, 35*len(g['vars'])+80))
                if img:
                    tmp_imgs.append(img)
                    doc.add_picture(img, width=Inches(6))
                    doc.add_paragraph(f'Figure : {g["parent"]}', style='Caption')

                _add_comment_word(doc,
                                  auto_comment_binary_group({'parent': g['parent'], **r}),
                                  user_comments.get(f'descriptive_bin:{g["parent"]}', ''))
            except Exception as e:
                doc.add_paragraph(f'Erreur : {e}')
            doc.add_paragraph()
        section += 1
        doc.add_page_break()

    # ── 2. Statistiques descriptives catégorielles ───────────────────────────
    if 'descriptive' in selected_analyses and cat_vars:
        doc.add_heading(f'{section}. Statistiques descriptives — Variables catégorielles', level=1)
        for var in cat_vars:
            if var not in df.columns:
                continue
            doc.add_heading(f'Tri à plat — {var}', level=2)
            try:
                r = tris_a_plat(df, var)
                doc.add_paragraph(
                    f'N valides = {r["n_valid"]} | Manquantes : {r["missing"]} ({r["missing_pct"]}%)')

                tbl_df = pd.read_html(StringIO(r['table_html']))[0]

                # Add table
                table = doc.add_table(rows=len(tbl_df)+1, cols=len(tbl_df.columns))
                table.style = 'Table Grid'

                # Header
                hdr_cells = table.rows[0].cells
                for j, col in enumerate(tbl_df.columns):
                    hdr_cells[j].text = str(col)
                    hdr_cells[j].paragraphs[0].runs[0].font.bold = True

                # Data
                for i, row in enumerate(tbl_df.values):
                    row_cells = table.rows[i+1].cells
                    for j, val in enumerate(row):
                        row_cells[j].text = str(val)

                doc.add_paragraph()

                # Add chart
                img = _fig_to_img(r['chart_bar'], width=520, height=280)
                if img:
                    tmp_imgs.append(img)
                    doc.add_picture(img, width=Inches(6))
                    doc.add_paragraph(f'Figure : Distribution de {var}', style='Caption')

                _add_comment_word(doc,
                                  auto_comment_categorical(r),
                                  user_comments.get(f'descriptive_cat:{var}', ''))
            except Exception as e:
                doc.add_paragraph(f'Erreur : {e}')
            doc.add_paragraph()
        section += 1
        doc.add_page_break()

    # ── 3. Statistiques descriptives continues ───────────────────────────────
    if 'descriptive' in selected_analyses and num_vars:
        doc.add_heading(f'{section}. Statistiques descriptives — Variables continues', level=1)
        for var in num_vars:
            if var not in df.columns:
                continue
            doc.add_heading(f'Distribution — {var}', level=2)
            try:
                r = statistiques_continues(df, var)
                tbl_df = pd.read_html(StringIO(r['table_html']))[0]

                # Add table
                table = doc.add_table(rows=len(tbl_df)+1, cols=len(tbl_df.columns))
                table.style = 'Table Grid'

                # Header
                hdr_cells = table.rows[0].cells
                for j, col in enumerate(tbl_df.columns):
                    hdr_cells[j].text = str(col)
                    hdr_cells[j].paragraphs[0].runs[0].font.bold = True

                # Data
                for i, row in enumerate(tbl_df.values):
                    row_cells = table.rows[i+1].cells
                    for j, val in enumerate(row):
                        row_cells[j].text = str(val)

                doc.add_paragraph()

                # Add chart
                img = _fig_to_img(r['chart_hist'], width=520, height=300)
                if img:
                    tmp_imgs.append(img)
                    doc.add_picture(img, width=Inches(6))
                    doc.add_paragraph(f'Figure : Distribution de {var}', style='Caption')

                _add_comment_word(doc,
                                  auto_comment_continuous(r),
                                  user_comments.get(f'descriptive_num:{var}', ''))
            except Exception as e:
                doc.add_paragraph(f'Erreur : {e}')
            doc.add_paragraph()
        section += 1
        doc.add_page_break()

    # ── 4. Tableaux croisés ──────────────────────────────────────────────────
    all_cat = cat_vars or [c for c, t in var_types.items() if t == 'categorielle']
    if 'crosstabs' in selected_analyses and len(all_cat) >= 2:
        doc.add_heading(f'{section}. Tableaux croisés', level=1)
        row_var = crosstabs_row if crosstabs_row in all_cat else all_cat[0]
        col_var = crosstabs_col if crosstabs_col in all_cat else (all_cat[1] if len(all_cat) > 1 else all_cat[0])
        if row_var == col_var and len(all_cat) > 1:
            col_var = all_cat[1] if all_cat[0] == row_var else all_cat[0]

        try:
            r = tableau_croise(df, row_var, col_var, 'row')
            doc.add_paragraph(f'Variables : {row_var} × {col_var}')
            doc.add_paragraph(
                f'χ² = {r["chi2"]} | p-value = {r["p_val"]} | '
                f'V de Cramér = {r["cramers_v"]} — '
                f'Liaison {r["assoc_label"]} ({r["sig_label"]})')

            ct = pd.crosstab(df[row_var].dropna().astype(str),
                              df[col_var].dropna().astype(str))
            ct_reset = ct.reset_index()

            # Add table
            table = doc.add_table(rows=len(ct_reset)+1, cols=len(ct_reset.columns))
            table.style = 'Table Grid'

            # Header
            hdr_cells = table.rows[0].cells
            for j, col in enumerate(ct_reset.columns):
                hdr_cells[j].text = str(col)
                hdr_cells[j].paragraphs[0].runs[0].font.bold = True

            # Data
            for i, row in enumerate(ct_reset.values):
                row_cells = table.rows[i+1].cells
                for j, val in enumerate(row):
                    row_cells[j].text = str(val)

            # Add chart
            img = _fig_to_img(r['chart_bar'], width=520, height=320)
            if img:
                tmp_imgs.append(img)
                doc.add_picture(img, width=Inches(6))

            _add_comment_word(doc,
                              auto_comment_crosstab(r),
                              user_comments.get(f'crosstabs:{row_var}|{col_var}', ''))
        except Exception as e:
            doc.add_paragraph(f'Erreur : {e}')
        section += 1
        doc.add_page_break()

    # ── 5. Analyse brute (overview & data quality) ───────────────────────────
    if 'raw' in selected_analyses:
        doc.add_heading(f'{section}. Analyse brute', level=1)

        def _docx_df_table(df_in: pd.DataFrame, max_rows: int = 80):
            d = df_in.head(max_rows)
            t = doc.add_table(rows=len(d)+1, cols=len(d.columns))
            t.style = 'Table Grid'
            hdr = t.rows[0].cells
            for j, c in enumerate(d.columns):
                hdr[j].text = str(c)
                try:
                    hdr[j].paragraphs[0].runs[0].font.bold = True
                except Exception:
                    pass
            for i, row in enumerate(d.values):
                cells = t.rows[i+1].cells
                for j, v in enumerate(row):
                    cells[j].text = '' if pd.isna(v) else str(v)

        try:
            sysa = systematic_analysis(df)
            desc = descriptive_summary(df)
            ov   = overview_stats(df)

            # ── 5.1 Qualité & Profil ─────────────────────────────────────────
            doc.add_heading(f'{section}.1 Qualité & Profil', level=2)
            q_img = _fig_to_img(data_quality_gauge(sysa.get('quality', {})), width=420, height=240)
            if q_img:
                tmp_imgs.append(q_img)
                doc.add_picture(q_img, width=Inches(5.5))
                doc.add_paragraph('Figure : Score de qualité des données')
            comp_img = _fig_to_img(composition_chart(sysa), width=420, height=240)
            if comp_img:
                tmp_imgs.append(comp_img)
                doc.add_picture(comp_img, width=Inches(5.5))
                doc.add_paragraph('Figure : Composition des variables')
            prof = sysa.get('profile', {}) or {}
            if prof:
                prof_df = pd.DataFrame([
                    ['Observations', prof.get('n_obs', '—')],
                    ['Variables', prof.get('n_vars', '—')],
                    ['Complètes', prof.get('complete_vars', '—')],
                    ['Avec manquantes', prof.get('vars_with_missing', '—')],
                    ['Continues', prof.get('n_continuous', '—')],
                    ['Catégorielles', prof.get('n_categorical', '—')],
                    ['Binaires', prof.get('n_binary', '—')],
                ], columns=['Indicateur', 'Valeur'])
                _docx_df_table(prof_df)

            # ── 5.2 Variables continues ──────────────────────────────────────
            doc.add_page_break()
            doc.add_heading(f'{section}.2 Variables continues', level=2)
            cont_list = desc.get('continuous', []) or []
            if not cont_list:
                doc.add_paragraph('Aucune variable continue détectée.')
            for r in cont_list[:10]:
                doc.add_heading(r['var'], level=3)
                stats_df = pd.DataFrame(list(r['stats'].items()), columns=['Statistique', 'Valeur'])
                _docx_df_table(stats_df)
                h_img = _fig_to_img(r.get('chart_hist'), width=480, height=260)
                if h_img:
                    tmp_imgs.append(h_img); doc.add_picture(h_img, width=Inches(6))
                b_img = _fig_to_img(r.get('chart_box'), width=480, height=260)
                if b_img:
                    tmp_imgs.append(b_img); doc.add_picture(b_img, width=Inches(6))

            # ── 5.3 Catégorielles ────────────────────────────────────────────
            doc.add_page_break()
            doc.add_heading(f'{section}.3 Catégorielles', level=2)
            cat_list = desc.get('categorical', []) or []
            if not cat_list:
                doc.add_paragraph('Aucune variable catégorielle détectée.')
            for r in cat_list[:10]:
                doc.add_heading(r['variable'], level=3)
                doc.add_paragraph(
                    f"{r['n_categories']} catégorie(s) — dominante : "
                    f"{r['dominant']} ({r['dominant_pct']}%)"
                )
                try:
                    freq_df = pd.read_html(StringIO(r['table']))[0]
                    _docx_df_table(freq_df)
                except Exception:
                    pass
                bar_img = _fig_to_img(r.get('chart_bar'), width=480, height=300)
                if bar_img:
                    tmp_imgs.append(bar_img); doc.add_picture(bar_img, width=Inches(6))

            # ── 5.4 Binaires ─────────────────────────────────────────────────
            doc.add_page_break()
            doc.add_heading(f'{section}.4 Binaires', level=2)
            bin_list = desc.get('binary', []) or []
            if not bin_list:
                doc.add_paragraph('Aucune variable binaire détectée.')
            else:
                bin_df = pd.DataFrame([[r['variable'], r['n_valid'], r['n_missing'],
                                          f"{r['pct_missing']}%", r['label_pos'],
                                          f"{r['prevalence']}%"] for r in bin_list],
                                       columns=['Variable', 'N valides', 'Manquantes',
                                                 '% manq.', 'Modalité positive', 'Prévalence'])
                _docx_df_table(bin_df)

            # ── 5.5 Manquantes ───────────────────────────────────────────────
            doc.add_page_break()
            doc.add_heading(f'{section}.5 Manquantes', level=2)
            try:
                mb = missing_bar_chart(df)
                if mb:
                    b_img = _fig_to_img(mb, width=520, height=300)
                    if b_img:
                        tmp_imgs.append(b_img); doc.add_picture(b_img, width=Inches(6.5))
                        doc.add_paragraph('Figure : Taux de données manquantes par variable')
            except Exception:
                pass
            try:
                mh = missing_heatmap(df)
                if mh:
                    h_img = _fig_to_img(mh, width=520, height=300)
                    if h_img:
                        tmp_imgs.append(h_img); doc.add_picture(h_img, width=Inches(6.5))
                        doc.add_paragraph('Figure : Carte des données manquantes')
            except Exception:
                pass

            # ── 5.6 Corrélations ─────────────────────────────────────────────
            doc.add_page_break()
            doc.add_heading(f'{section}.6 Corrélations', level=2)
            try:
                corr = correlation_matrix(df)
                if corr:
                    c_img = _fig_to_img(corr, width=520, height=400)
                    if c_img:
                        tmp_imgs.append(c_img); doc.add_picture(c_img, width=Inches(6))
                        doc.add_paragraph('Figure : Matrice de corrélations (variables numériques)')
                    else:
                        doc.add_paragraph('Pas assez de variables numériques pour une matrice de corrélations.')
                else:
                    doc.add_paragraph('Pas assez de variables numériques pour une matrice de corrélations.')
            except Exception:
                doc.add_paragraph('Pas assez de variables numériques pour une matrice de corrélations.')

            # ── 5.7 Vue d'ensemble ───────────────────────────────────────────
            doc.add_page_break()
            doc.add_heading(f"{section}.7 Vue d'ensemble", level=2)
            try:
                ov_df = pd.read_html(StringIO(ov['table_html']))[0]
                _docx_df_table(ov_df, max_rows=200)
            except Exception as e:
                doc.add_paragraph(f"Vue d'ensemble indisponible : {e}")

            # ── Annexe : anomalies / recommandations ─────────────────────────
            if sysa.get('anomalies') or sysa.get('recommendations'):
                doc.add_heading(f'{section}.8 Anomalies & Recommandations', level=2)
                if sysa.get('anomalies'):
                    for a in sysa['anomalies'][:20]:
                        doc.add_paragraph(f"- {a['type']} — {a['variable']} : {a['note']}")
                if sysa.get('recommendations'):
                    doc.add_paragraph('Recommandations :').runs[0].font.bold = True
                    for r in sysa['recommendations'][:10]:
                        doc.add_paragraph(f"• {r['type']} — {r['rationale']}")
        except Exception as e:
            doc.add_paragraph(f'Erreur lors de l\'analyse brute : {e}')
        section += 1
        doc.add_page_break()

    # ── 6. Multi-enquête (DPF / PHAKTS) ──────────────────────────────────────
    if 'multi_survey' in selected_analyses and multi_survey_results:
        doc.add_heading(f'{section}. Analyse multi-enquête (DPF / PHAKTS)', level=1)
        if multi_survey_meta:
            doc.add_paragraph(
                f"{multi_survey_meta['n_surveys']} enquête(s) fusionnée(s) selon le mode "
                f"« {multi_survey_meta['mode']} » — {multi_survey_meta['n_obs']} observations, "
                f"{multi_survey_meta['n_vars']} variables alignées."
            )
        for idx, r in enumerate(multi_survey_results, start=1):
            doc.add_heading(f"{section}.{idx} Variable « {r.get('variable', '?')} »", level=2)
            if r.get('error'):
                doc.add_paragraph(f"Erreur : {r['error']}")
                continue
            try:
                if r.get('kind') == 'continuous' and r.get('summary_html'):
                    doc.add_paragraph('Résumé statistique par enquête :')
                    summ_df = pd.read_html(StringIO(r['summary_html']))[0]
                    _docx_df_table(summ_df)
                elif r.get('kind') == 'categorical':
                    if r.get('count_html'):
                        doc.add_paragraph('Effectifs :')
                        c_df = pd.read_html(StringIO(r['count_html']))[0]
                        _docx_df_table(c_df)
                    if r.get('pct_html'):
                        doc.add_paragraph('Pourcentages (% colonne) :')
                        p_df = pd.read_html(StringIO(r['pct_html']))[0]
                        _docx_df_table(p_df)
            except Exception as e:
                doc.add_paragraph(f'Tableaux indisponibles : {e}')
            # Graphique
            chart_img = _fig_to_img(r.get('chart'), width=520, height=320)
            if chart_img:
                tmp_imgs.append(chart_img)
                doc.add_picture(chart_img, width=Inches(6.5))
                doc.add_paragraph(
                    f"Figure : Comparaison de « {r.get('variable', '?')} » entre enquêtes"
                )
        section += 1
        doc.add_page_break()

    # ── Conclusion ───────────────────────────────────────────────────────────
    doc.add_heading('Conclusion', level=1)
    conclusion = doc.add_paragraph(
        f'Ce rapport a été généré automatiquement le '
        f'{datetime.now().strftime("%d/%m/%Y")} par SPAD Analyzer. '
        f'Il couvre {df.shape[0]} observations et '
        f'{df.shape[1]} variables.')

    if output_path:
        doc.save(output_path)
    else:
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    # Clean up temp images
    for p in tmp_imgs:
        try:
            os.remove(p)
        except Exception:
            pass

    return None


def generate_excel_report(df: pd.DataFrame, output_path: str = None,
                           title: str = 'Rapport', author: str = 'Analyste',
                           selected_analyses: list = None,
                           cat_vars: list = None, num_vars: list = None,
                           bin_groups: list = None,
                           crosstabs_row=None,
                           crosstabs_col=None,
                           user_comments: dict = None,
                           multi_survey_results: list = None,
                           multi_survey_meta: dict = None):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter
    from openpyxl.drawing.image import Image as XLImage

    tmp_imgs = []
    bin_groups = bin_groups or []
    user_comments = user_comments or {}
    if selected_analyses is None:
        selected_analyses = ['raw']
    cat_vars = cat_vars or []
    num_vars = num_vars or []
    var_types = get_var_types(df)

    HEADER_FILL  = PatternFill('solid', fgColor='1A3C5E')
    HEADER_FONT  = Font(color='FFFFFF', bold=True, size=10)
    TITLE_FONT   = Font(color='1A3C5E', bold=True, size=14)
    SUBHEAD_FONT = Font(color='E05C1A', bold=True, size=11)
    BODY_FONT    = Font(size=10)
    ITALIC_FONT  = Font(size=10, italic=True)
    BOLD_FONT    = Font(size=10, bold=True)

    wb = Workbook()
    ws0 = wb.active
    ws0.title = 'Résumé'

    def _sheet(name):
        safe = re.sub(r'[\[\]\:\*\?/\\]', '', name)[:31] or 'Feuille'
        base, i = safe, 1
        while safe in wb.sheetnames:
            i += 1
            suffix = f'_{i}'
            safe = base[:31 - len(suffix)] + suffix
        return wb.create_sheet(safe)

    def _write_title(ws, text, row=1):
        ws.cell(row=row, column=1, value=text).font = TITLE_FONT
        return row + 2

    def _write_subheading(ws, text, row):
        ws.cell(row=row, column=1, value=text).font = SUBHEAD_FONT
        return row + 1

    def _write_note(ws, text, row, italic=False):
        ws.cell(row=row, column=1, value=text).font = ITALIC_FONT if italic else BODY_FONT
        return row + 1

    def _write_df(ws, df_in, row, max_rows=200):
        d = df_in.head(max_rows)
        for j, col in enumerate(d.columns, start=1):
            cell = ws.cell(row=row, column=j, value=str(col))
            cell.font = HEADER_FONT
            cell.fill = HEADER_FILL
            cell.alignment = Alignment(horizontal='center')
        for i, vals in enumerate(d.values, start=1):
            for j, v in enumerate(vals, start=1):
                v_out = '' if (isinstance(v, float) and pd.isna(v)) else v
                ws.cell(row=row + i, column=j, value=v_out).font = BODY_FONT
        for j, col in enumerate(d.columns, start=1):
            try:
                sample = [len(str(v)) for v in d[col].head(50)]
            except Exception:
                sample = []
            width = max(10, min(40, max([len(str(col))] + sample) + 2))
            ws.column_dimensions[get_column_letter(j)].width = width
        return row + len(d) + 2

    def _write_image(ws, img_path, row, width_px=520):
        if not img_path:
            return row
        try:
            img = XLImage(img_path)
            scale = width_px / float(img.width)
            img.width = width_px
            img.height = int(img.height * scale)
            ws.add_image(img, f'A{row}')
            rows_used = max(2, int(img.height / 15) + 1)
            return row + rows_used
        except Exception:
            return row

    # ── Résumé ────────────────────────────────────────────────────────────
    row = _write_title(ws0, title)
    row = _write_note(ws0, "Rapport d'analyse statistique — SPAD Analyzer", row)
    row += 1
    for k, v in [
        ['Auteur', author],
        ['Date', datetime.now().strftime('%d/%m/%Y à %H:%M')],
        ['Observations', df.shape[0]],
        ['Variables analysées', df.shape[1]],
        ['Données manquantes', int(df.isnull().sum().sum())],
    ]:
        ws0.cell(row=row, column=1, value=k).font = BOLD_FONT
        ws0.cell(row=row, column=2, value=v).font = BODY_FONT
        row += 1
    ws0.column_dimensions['A'].width = 28
    ws0.column_dimensions['B'].width = 40

    section = 1

    # ── 1. Questions à choix multiples ──────────────────────────────────────
    if 'binaire' in selected_analyses and bin_groups:
        ws = _sheet('Choix multiples')
        row = _write_title(ws, f'{section}. Questions à choix multiples')
        for g in bin_groups:
            row = _write_subheading(ws, g['parent'], row)
            try:
                r = analyse_binaire_groupe(df, g['vars'])
                tbl_df = pd.read_html(StringIO(r['table_html']))[0]
                row = _write_df(ws, tbl_df, row)
                img = _fig_to_img(r['chart'], width=520, height=max(250, 35 * len(g['vars']) + 80))
                if img:
                    tmp_imgs.append(img)
                    row = _write_image(ws, img, row)
                auto = auto_comment_binary_group({'parent': g['parent'], **r})
                user = user_comments.get(f'descriptive_bin:{g["parent"]}', '')
                if auto:
                    row = _write_note(ws, f'Lecture automatique : {auto}', row)
                if user:
                    row = _write_note(ws, f'Commentaire de l\'analyste : {user}', row, italic=True)
            except Exception as e:
                row = _write_note(ws, f'Erreur : {e}', row)
            row += 1
        section += 1

    # ── 2. Statistiques descriptives catégorielles ──────────────────────────
    if 'descriptive' in selected_analyses and cat_vars:
        ws = _sheet('Descriptif catégoriel')
        row = _write_title(ws, f'{section}. Statistiques descriptives — Variables catégorielles')
        for var in cat_vars:
            if var not in df.columns:
                continue
            row = _write_subheading(ws, f'Tri à plat — {var}', row)
            try:
                r = tris_a_plat(df, var)
                row = _write_note(ws, f'N valides = {r["n_valid"]} | Manquantes : {r["missing"]} ({r["missing_pct"]}%)', row)
                tbl_df = pd.read_html(StringIO(r['table_html']))[0]
                row = _write_df(ws, tbl_df, row)
                img = _fig_to_img(r['chart_bar'], width=520, height=280)
                if img:
                    tmp_imgs.append(img)
                    row = _write_image(ws, img, row)
                auto = auto_comment_categorical(r)
                user = user_comments.get(f'descriptive_cat:{var}', '')
                if auto:
                    row = _write_note(ws, f'Lecture automatique : {auto}', row)
                if user:
                    row = _write_note(ws, f'Commentaire de l\'analyste : {user}', row, italic=True)
            except Exception as e:
                row = _write_note(ws, f'Erreur : {e}', row)
            row += 1
        section += 1

    # ── 3. Statistiques descriptives continues ──────────────────────────────
    if 'descriptive' in selected_analyses and num_vars:
        ws = _sheet('Descriptif continues')
        row = _write_title(ws, f'{section}. Statistiques descriptives — Variables continues')
        for var in num_vars:
            if var not in df.columns:
                continue
            row = _write_subheading(ws, f'Distribution — {var}', row)
            try:
                r = statistiques_continues(df, var)
                tbl_df = pd.read_html(StringIO(r['table_html']))[0]
                row = _write_df(ws, tbl_df, row)
                img = _fig_to_img(r['chart_hist'], width=520, height=300)
                if img:
                    tmp_imgs.append(img)
                    row = _write_image(ws, img, row)
                auto = auto_comment_continuous(r)
                user = user_comments.get(f'descriptive_num:{var}', '')
                if auto:
                    row = _write_note(ws, f'Lecture automatique : {auto}', row)
                if user:
                    row = _write_note(ws, f'Commentaire de l\'analyste : {user}', row, italic=True)
            except Exception as e:
                row = _write_note(ws, f'Erreur : {e}', row)
            row += 1
        section += 1

    # ── 4. Tableaux croisés ──────────────────────────────────────────────────
    all_cat = cat_vars or [c for c, t in var_types.items() if t == 'categorielle']
    if 'crosstabs' in selected_analyses and len(all_cat) >= 2:
        ws = _sheet('Tableaux croisés')
        row = _write_title(ws, f'{section}. Tableaux croisés')
        row_var = crosstabs_row if crosstabs_row in all_cat else all_cat[0]
        col_var = crosstabs_col if crosstabs_col in all_cat else (all_cat[1] if len(all_cat) > 1 else all_cat[0])
        if row_var == col_var and len(all_cat) > 1:
            col_var = all_cat[1] if all_cat[0] == row_var else all_cat[0]
        try:
            r = tableau_croise(df, row_var, col_var, 'row')
            row = _write_note(ws, f'Variables : {row_var} × {col_var}', row)
            row = _write_note(ws, (
                f'χ² = {r["chi2"]} | p-value = {r["p_val"]} | '
                f'V de Cramér = {r["cramers_v"]} — '
                f'Liaison {r["assoc_label"]} ({r["sig_label"]})'), row)
            ct = pd.crosstab(df[row_var].dropna().astype(str), df[col_var].dropna().astype(str))
            row = _write_df(ws, ct.reset_index(), row)
            img = _fig_to_img(r['chart_bar'], width=520, height=320)
            if img:
                tmp_imgs.append(img)
                row = _write_image(ws, img, row)
            auto = auto_comment_crosstab(r)
            user = user_comments.get(f'crosstabs:{row_var}|{col_var}', '')
            if auto:
                row = _write_note(ws, f'Lecture automatique : {auto}', row)
            if user:
                row = _write_note(ws, f'Commentaire de l\'analyste : {user}', row, italic=True)
        except Exception as e:
            row = _write_note(ws, f'Erreur : {e}', row)
        section += 1

    # ── 5. Analyse brute ─────────────────────────────────────────────────────
    if 'raw' in selected_analyses:
        try:
            sysa = systematic_analysis(df)
            desc = descriptive_summary(df)
            ov   = overview_stats(df)

            # 5.1 Qualité & Profil
            ws = _sheet('Brute-Qualité')
            row = _write_title(ws, f'{section}.1 Qualité & Profil')
            q_img = _fig_to_img(data_quality_gauge(sysa.get('quality', {})), width=420, height=240)
            if q_img:
                tmp_imgs.append(q_img)
                row = _write_image(ws, q_img, row)
            comp_img = _fig_to_img(composition_chart(sysa), width=420, height=240)
            if comp_img:
                tmp_imgs.append(comp_img)
                row = _write_image(ws, comp_img, row)
            prof = sysa.get('profile', {}) or {}
            if prof:
                prof_df = pd.DataFrame([
                    ['Observations', prof.get('n_obs', '—')],
                    ['Variables', prof.get('n_vars', '—')],
                    ['Complètes', prof.get('complete_vars', '—')],
                    ['Avec manquantes', prof.get('vars_with_missing', '—')],
                    ['Continues', prof.get('n_continuous', '—')],
                    ['Catégorielles', prof.get('n_categorical', '—')],
                    ['Binaires', prof.get('n_binary', '—')],
                ], columns=['Indicateur', 'Valeur'])
                row = _write_df(ws, prof_df, row)

            # 5.2 Variables continues
            ws = _sheet('Brute-Continues')
            row = _write_title(ws, f'{section}.2 Variables continues')
            cont_list = desc.get('continuous', []) or []
            if not cont_list:
                row = _write_note(ws, 'Aucune variable continue détectée.', row)
            for r in cont_list[:10]:
                row = _write_subheading(ws, r['var'], row)
                stats_df = pd.DataFrame(list(r['stats'].items()), columns=['Statistique', 'Valeur'])
                row = _write_df(ws, stats_df, row)
                h_img = _fig_to_img(r.get('chart_hist'), width=480, height=260)
                if h_img:
                    tmp_imgs.append(h_img); row = _write_image(ws, h_img, row)
                b_img = _fig_to_img(r.get('chart_box'), width=480, height=260)
                if b_img:
                    tmp_imgs.append(b_img); row = _write_image(ws, b_img, row)

            # 5.3 Catégorielles
            ws = _sheet('Brute-Catégorielles')
            row = _write_title(ws, f'{section}.3 Catégorielles')
            cat_list = desc.get('categorical', []) or []
            if not cat_list:
                row = _write_note(ws, 'Aucune variable catégorielle détectée.', row)
            for r in cat_list[:10]:
                row = _write_subheading(ws, r['variable'], row)
                row = _write_note(ws, f"{r['n_categories']} catégorie(s) — dominante : {r['dominant']} ({r['dominant_pct']}%)", row)
                try:
                    freq_df = pd.read_html(StringIO(r['table']))[0]
                    row = _write_df(ws, freq_df, row)
                except Exception:
                    pass
                bar_img = _fig_to_img(r.get('chart_bar'), width=480, height=300)
                if bar_img:
                    tmp_imgs.append(bar_img); row = _write_image(ws, bar_img, row)

            # 5.4 Binaires
            ws = _sheet('Brute-Binaires')
            row = _write_title(ws, f'{section}.4 Binaires')
            bin_list = desc.get('binary', []) or []
            if not bin_list:
                row = _write_note(ws, 'Aucune variable binaire détectée.', row)
            else:
                bin_df = pd.DataFrame([[r['variable'], r['n_valid'], r['n_missing'],
                                          f"{r['pct_missing']}%", r['label_pos'],
                                          f"{r['prevalence']}%"] for r in bin_list],
                                       columns=['Variable', 'N valides', 'Manquantes',
                                                 '% manq.', 'Modalité positive', 'Prévalence'])
                row = _write_df(ws, bin_df, row)

            # 5.5 Manquantes
            ws = _sheet('Brute-Manquantes')
            row = _write_title(ws, f'{section}.5 Manquantes')
            try:
                mb = missing_bar_chart(df)
                if mb:
                    b_img = _fig_to_img(mb, width=520, height=300)
                    if b_img:
                        tmp_imgs.append(b_img); row = _write_image(ws, b_img, row)
            except Exception:
                pass
            try:
                mh = missing_heatmap(df)
                if mh:
                    h_img = _fig_to_img(mh, width=520, height=300)
                    if h_img:
                        tmp_imgs.append(h_img); row = _write_image(ws, h_img, row)
            except Exception:
                pass

            # 5.6 Corrélations
            ws = _sheet('Brute-Corrélations')
            row = _write_title(ws, f'{section}.6 Corrélations')
            try:
                corr = correlation_matrix(df)
                c_img = _fig_to_img(corr, width=520, height=400) if corr else None
                if c_img:
                    tmp_imgs.append(c_img); row = _write_image(ws, c_img, row)
                else:
                    row = _write_note(ws, 'Pas assez de variables numériques pour une matrice de corrélations.', row)
            except Exception:
                row = _write_note(ws, 'Pas assez de variables numériques pour une matrice de corrélations.', row)

            # 5.7 Vue d'ensemble
            ws = _sheet("Brute-Ensemble")
            row = _write_title(ws, f"{section}.7 Vue d'ensemble")
            try:
                ov_df = pd.read_html(StringIO(ov['table_html']))[0]
                row = _write_df(ws, ov_df, row, max_rows=500)
            except Exception as e:
                row = _write_note(ws, f"Vue d'ensemble indisponible : {e}", row)

            # 5.8 Anomalies & recommandations
            if sysa.get('anomalies') or sysa.get('recommendations'):
                ws = _sheet('Brute-Anomalies')
                row = _write_title(ws, f'{section}.8 Anomalies & Recommandations')
                if sysa.get('anomalies'):
                    for a in sysa['anomalies'][:20]:
                        row = _write_note(ws, f"- {a['type']} — {a['variable']} : {a['note']}", row)
                if sysa.get('recommendations'):
                    row = _write_note(ws, 'Recommandations :', row)
                    for r in sysa['recommendations'][:10]:
                        row = _write_note(ws, f"• {r['type']} — {r['rationale']}", row)
        except Exception as e:
            ws = _sheet('Brute-Erreur')
            _write_note(ws, f"Erreur lors de l'analyse brute : {e}", 1)
        section += 1

    # ── 6. Multi-enquête (DPF / PHAKTS) ──────────────────────────────────────
    if 'multi_survey' in selected_analyses and multi_survey_results:
        ws = _sheet('Multi-enquête')
        row = _write_title(ws, f'{section}. Analyse multi-enquête (DPF / PHAKTS)')
        if multi_survey_meta:
            row = _write_note(ws, (
                f"{multi_survey_meta['n_surveys']} enquête(s) fusionnée(s) selon le mode "
                f"« {multi_survey_meta['mode']} » — {multi_survey_meta['n_obs']} observations, "
                f"{multi_survey_meta['n_vars']} variables alignées."), row)
        for idx, r in enumerate(multi_survey_results, start=1):
            row = _write_subheading(ws, f"{section}.{idx} Variable « {r.get('variable', '?')} »", row)
            if r.get('error'):
                row = _write_note(ws, f"Erreur : {r['error']}", row)
                continue
            try:
                if r.get('kind') == 'continuous' and r.get('summary_html'):
                    row = _write_note(ws, 'Résumé statistique par enquête :', row)
                    summ_df = pd.read_html(StringIO(r['summary_html']))[0]
                    row = _write_df(ws, summ_df, row)
                elif r.get('kind') == 'categorical':
                    if r.get('count_html'):
                        row = _write_note(ws, 'Effectifs :', row)
                        c_df = pd.read_html(StringIO(r['count_html']))[0]
                        row = _write_df(ws, c_df, row)
                    if r.get('pct_html'):
                        row = _write_note(ws, 'Pourcentages (% colonne) :', row)
                        p_df = pd.read_html(StringIO(r['pct_html']))[0]
                        row = _write_df(ws, p_df, row)
            except Exception as e:
                row = _write_note(ws, f'Tableaux indisponibles : {e}', row)
            chart_img = _fig_to_img(r.get('chart'), width=520, height=320)
            if chart_img:
                tmp_imgs.append(chart_img)
                row = _write_image(ws, chart_img, row)
        section += 1

    if output_path:
        wb.save(output_path)
    else:
        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        for p in tmp_imgs:
            try:
                os.remove(p)
            except Exception:
                pass
        return buf

    for p in tmp_imgs:
        try:
            os.remove(p)
        except Exception:
            pass

    return None
