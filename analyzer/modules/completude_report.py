"""
SPAD Analyzer — Rapport de complétude au format Word (.docx)

Reproduit la structure d'un compte rendu de débriefing terrain type
(indicateurs clés / points d'alerte / complétude par formulaire / complétude
par district / actions), à partir des données déjà calculées par le moteur
de complétude (modules/completeness.py) — voir le modèle fourni par
l'utilisateur : « Compte rendu Debriefing_UGP » (SPAD 2026, pilote).

Seul ce qui est objectivement dérivable des données Kobo est généré
automatiquement (taux, anomalies détectées). Les décisions de réunion
(budget, chronogramme, attribution nominative des actions) ne le sont pas
— elles n'existent pas dans les données et seraient inventées si on les
générait.
"""
import io
import os
from datetime import datetime, date

import plotly.graph_objects as go

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from modules import reference_data as rd
from modules import completeness as cp
from modules import tendance
# L'import déclenche _setup_kaleido() (contournement du bug Kaleido avec les
# chemins contenant des espaces — voir modules/report_generator.py) et fournit
# _fig_to_img(), déjà utilisé pour les graphiques du rapport PDF/Word existant.
from modules.report_generator import _fig_to_img

STATUT_COLOR = {
    'zero': '#A0AEC0', 'en_cours': '#E67E22', 'cible': '#2E7D4F',
    'verifier': '#C0392B', 'suivi': '#A0AEC0', 'inconnu': '#D9DEE4',
}

PRIMARY = RGBColor(26, 60, 94)
ACCENT  = RGBColor(224, 92, 26)

STATUT_LABEL = {
    'zero': '0 %', 'en_cours': 'En cours', 'cible': 'Cible atteinte',
    'verifier': 'À vérifier', 'suivi': 'Suivi', 'inconnu': 'Non calculé',
}



def _shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def _style_header_row(row, hex_color='1A3C5E'):
    for cell in row.cells:
        _shade_cell(cell, hex_color)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(9)


def _add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    _style_header_row(table.rows[0])
    for row_vals in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_vals):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table


def _national_totals(national, form_codes):
    """Taux global agrégé (somme reçu / somme cible) sur les formulaires
    effectivement mappés (statut != 'inconnu')."""
    recu_total, cible_total, n_actifs = 0, 0, 0
    for code in form_codes:
        r = national.get(code, {})
        if r.get('statut') == 'inconnu':
            continue
        n_actifs += 1
        recu_total += r.get('recu') or 0
        cible_total += r.get('cible') or 0
    taux_global = round(100 * recu_total / cible_total, 1) if cible_total else None
    return recu_total, cible_total, taux_global, n_actifs


def _district_form_bullets(district_table, form_codes):
    """Un commentaire PAR FORMULAIRE (« par enquête »), pas par volet
    agrégé : district le plus avancé et le moins avancé, avec le nombre de
    soumissions qu'il reste à collecter pour ce dernier — l'information la
    plus directement actionnable pour le suivi terrain, plutôt qu'une
    moyenne qui masque les extrêmes."""
    bullets = []
    for code in form_codes:
        entries = []
        for d in district_table.values():
            cell = d['forms'].get(code, {})
            if cell.get('statut') not in (None, 'inconnu'):
                entries.append((d['nom'], cell.get('taux') or 0, cell.get('recu', 0), cell.get('cible') or 0))
        if not entries:
            continue
        best = max(entries, key=lambda x: x[1])
        worst = min(entries, key=lambda x: x[1])
        gap = max(0, worst[3] - worst[2])
        bullets.append(
            f"{code} — {rd.FORM_LABELS[code]} : district le plus avancé = {best[0]} ({best[1]} %) · "
            f"le moins avancé = {worst[0]} ({worst[1]} %"
            + (f", {gap} soumission(s) restante(s) pour atteindre la cible)" if gap else ", cible déjà atteinte)")
        )
    return bullets


def _legend_right(**overrides):
    """Légende verticale, plaquée à l'extrême droite du graphique, bien
    encadrée (fond blanc + bordure) pour bien la détacher visuellement de
    la zone de traçage plutôt que de la laisser flotter au-dessus."""
    base = dict(
        orientation='v', x=1.02, xanchor='left', y=1, yanchor='top',
        bgcolor='white', bordercolor='#D9DEE4', borderwidth=1,
        font=dict(size=10),
    )
    base.update(overrides)
    return base


def _grouped_bar_chart(items, form_codes, title, width=830, height=380):
    """Barres groupées (une série par formulaire, une catégorie par unité)
    — permet de comparer plusieurs formulaires en un coup d'œil plutôt que
    de ne lire qu'un tableau. Utilisé pour district et superviseur (peu
    d'unités, noms lisibles en abscisse)."""
    noms = [v['nom'] for v in items.values()]
    if not noms:
        return None
    fig = go.Figure()
    for code in form_codes:
        y = [v['forms'].get(code, {}).get('taux') or 0 for v in items.values()]
        fig.add_trace(go.Bar(name=code, x=noms, y=y))
    fig.update_layout(
        title=dict(text=title, x=0.01, xanchor='left', y=0.98, yanchor='top', font=dict(size=14)),
        barmode='group', yaxis_title='Taux (%)',
        plot_bgcolor='white', paper_bgcolor='white',
        margin=dict(t=55, b=110, l=60, r=110), xaxis=dict(tickangle=-35),
        legend=_legend_right(),
        font=dict(size=11),
    )
    return _fig_to_img(fig.to_json(), width=width, height=height)


def _distribution_chart(items, form_codes, title, width=720, height=380):
    """Boîtes à moustaches : distribution des taux par formulaire — utilisé
    pour les enquêteurs (trop nombreux pour un graphique nominatif lisible),
    montre la dispersion et les valeurs atypiques d'un coup d'œil."""
    fig = go.Figure()
    any_data = False
    for code in form_codes:
        y = [v['forms'].get(code, {}).get('taux') for v in items.values()
             if v['forms'].get(code, {}).get('taux') is not None]
        if y:
            any_data = True
            fig.add_trace(go.Box(y=y, name=code, boxpoints='outliers'))
    if not any_data:
        return None
    fig.update_layout(
        title=dict(text=title, x=0.01, xanchor='left', y=0.97, yanchor='top', font=dict(size=14)),
        yaxis_title='Taux (%)',
        plot_bgcolor='white', paper_bgcolor='white',
        margin=dict(t=55, b=40, l=60, r=110),
        legend=_legend_right(),
        font=dict(size=11),
    )
    return _fig_to_img(fig.to_json(), width=width, height=height)


def _velocity_and_projection(national, form_codes, history):
    """Pour chaque formulaire, estime la vélocité de collecte (soumissions/
    jour) et projette le délai nécessaire pour atteindre la cible au rythme
    actuel — lecture de suivi-évaluation standard (comparer l'état à un
    instant T à sa dynamique récente, pas seulement au seuil final).

    Repose sur l'historique quotidien des TAUX (modules/tendance.py),
    reconverti en effectifs via la cible courante du formulaire — fixe dans
    le temps pour un même formulaire SPAD, donc sans perte d'information.
    Nécessite au moins 2 jours calendaires distincts avec un taux connu
    pour ce formulaire ; sinon l'estimation est explicitement signalée comme
    indisponible plutôt que devinée sur une seule observation."""
    out = {}
    for code in form_codes:
        r = national.get(code, {})
        cible, recu, statut = r.get('cible'), r.get('recu', 0), r.get('statut', 'inconnu')
        if not cible or statut == 'inconnu':
            out[code] = {'disponible': False}
            continue

        points = [(h['date'], h['taux'].get(code)) for h in history
                  if h.get('taux', {}).get(code) is not None]
        if len(points) < 2:
            out[code] = {'disponible': False}
            continue

        points = points[-8:]  # jusqu'à 7 intervalles (fenêtre ~1 semaine)
        d0, t0 = points[0]
        d1, t1 = points[-1]
        n_jours = (date.fromisoformat(d1) - date.fromisoformat(d0)).days
        if n_jours <= 0:
            out[code] = {'disponible': False}
            continue

        vitesse = (round(t1 / 100 * cible) - round(t0 / 100 * cible)) / n_jours
        entry = {'disponible': True, 'vitesse_jour': round(vitesse, 1), 'n_jours_observes': n_jours}
        if recu >= cible:
            entry['projection'] = 'cible_atteinte'
        elif vitesse > 0:
            entry['projection'] = 'estimee'
            entry['jours_restants'] = max(1, round((cible - recu) / vitesse))
        else:
            entry['projection'] = 'stagnant'
        out[code] = entry
    return out


def _summary_chart_image(national, form_codes):
    """Graphique de synthèse — taux de complétude par formulaire, coloré par
    statut (même palette que la page Graphiques de l'app). Résume en un coup
    d'œil le contenu détaillé du tableau de la section 3. Renvoie le chemin
    d'un fichier PNG temporaire (à supprimer par l'appelant), ou None si
    l'export a échoué (ex. Kaleido indisponible) — le rapport reste généré
    sans ce visuel plutôt que d'échouer entièrement."""
    labels, taux, colors, texts = [], [], [], []
    for code in form_codes:
        r = national.get(code, {})
        labels.append(code)
        t = r.get('taux')
        taux.append(t or 0)
        colors.append(STATUT_COLOR.get(r.get('statut', 'inconnu'), '#D9DEE4'))
        texts.append(f"{t} %" if t is not None else '—')

    fig = go.Figure(go.Bar(x=labels, y=taux, marker_color=colors, text=texts, textposition='outside'))
    fig.update_layout(
        margin=dict(t=15, r=15, b=40, l=50),
        yaxis=dict(title='Taux (%)', rangemode='tozero', range=[0, max([110] + [t + 20 for t in taux])]),
        xaxis=dict(title=None),
        plot_bgcolor='white', paper_bgcolor='white',
        font=dict(size=13),
    )
    return _fig_to_img(fig.to_json(), width=680, height=340)


def _form_counts_chart_image(national, form_codes):
    """Deuxième lecture du même tableau, en effectifs bruts (reçu vs cible)
    plutôt qu'en taux — un formulaire à 100 % peut représenter 1800
    soumissions (F5) ou 12 (F01) : le taux seul masque cet écart d'échelle,
    utile pour prioriser où porter l'effort de relance en pratique."""
    labels, recu_vals, cible_vals = [], [], []
    for code in form_codes:
        r = national.get(code, {})
        labels.append(code)
        recu_vals.append(r.get('recu') or 0)
        cible_vals.append(r.get('cible') or 0)
    if not labels:
        return None
    fig = go.Figure()
    fig.add_trace(go.Bar(name='Reçu', x=labels, y=recu_vals, marker_color='#1A3C5E',
                          text=recu_vals, textposition='outside'))
    fig.add_trace(go.Bar(name='Cible', x=labels, y=cible_vals, marker_color='#D9DEE4',
                          text=cible_vals, textposition='outside'))
    fig.update_layout(
        barmode='group', margin=dict(t=15, r=110, b=40, l=50),
        yaxis=dict(title='Soumissions'), xaxis=dict(title=None),
        plot_bgcolor='white', paper_bgcolor='white',
        legend=_legend_right(), font=dict(size=13),
    )
    return _fig_to_img(fig.to_json(), width=680, height=340)


def _form_rate_analysis_text(national, form_codes, anomalies_par_formulaire, district_table=None):
    """Analyse en prose du graphique des TAUX par formulaire — regroupe les
    formulaires par statut plutôt que de les énumérer un par un, pour une
    lecture synthétique cohérente avec ce que montre le graphique. Ajoute,
    quand la table par district est disponible, un rappel de la « Cible
    réelle atteinte » : un taux national ≥100 % peut cacher des districts
    encore en retard, compensés par d'autres très en avance."""
    atteints, en_cours, zero, verifier = [], [], [], []
    for code in form_codes:
        r = national.get(code, {})
        statut, taux = r.get('statut'), r.get('taux')
        label = f"{code} ({taux} %)" if taux is not None else code
        if statut == 'verifier':
            verifier.append(label)
        elif statut == 'cible':
            atteints.append(label)
        elif statut == 'en_cours':
            en_cours.append(label)
        elif statut == 'zero':
            zero.append(label)

    phrases = []
    if atteints or verifier:
        cibles = atteints + verifier
        phrases.append(f"{len(cibles)} formulaire(s) ont atteint leur cible : {', '.join(cibles)}.")
    if en_cours:
        phrases.append(f"{len(en_cours)} formulaire(s) restent en cours de collecte : {', '.join(en_cours)}.")
    if zero:
        phrases.append(f"{len(zero)} formulaire(s) affichent 0 % de soumission : {', '.join(zero)}.")
    if verifier:
        n_anom = sum(anomalies_par_formulaire.get(c.split(' ')[0], 0) for c in verifier)
        phrases.append(
            f"{len(verifier)} formulaire(s) dépassent largement leur cible (≥200 %) — "
            f"{', '.join(verifier)} — signe probable de doublons"
            + (f" ({n_anom} anomalie(s) déjà détectée(s) en excès sur ces formulaires)." if n_anom else ".")
        )
    if not phrases:
        return "Aucun formulaire mappé ne permet de calcul de taux à ce jour."

    if district_table:
        partiels, complets = [], []
        for code in form_codes:
            dr = cp.district_reel(district_table, code)
            if not dr:
                continue
            (complets if dr['pct'] >= 100 else partiels).append(
                f"{code} ({dr['atteints']}/{dr['total']} districts, {dr['pct']} %)"
            )
        if partiels:
            phrases.append(
                f"Cible réelle atteinte (proportion de districts individuellement ≥100 %, pas "
                f"seulement en moyenne nationale) : {', '.join(partiels)} n'y sont pas encore "
                f"partout" + (f", contrairement à {', '.join(complets)} (100 % des districts)." if complets else ".")
            )
        elif complets:
            phrases.append(
                "Cible réelle atteinte : tous les formulaires avec une cible réelle calculée "
                "l'ont atteinte dans l'ensemble de leurs districts."
            )
    return ' '.join(phrases)


def _form_counts_analysis_text(national, form_codes):
    """Analyse en prose du graphique reçu/cible en effectifs bruts — met en
    évidence l'écart le plus important en VOLUME (pas en %), l'information
    la plus actionnable pour prioriser une relance terrain."""
    manquants, excedents = [], []
    for code in form_codes:
        r = national.get(code, {})
        cible, recu = r.get('cible'), r.get('recu') or 0
        if not cible:
            continue
        ecart = recu - cible
        if ecart < 0:
            manquants.append((code, -ecart, recu, cible))
        elif ecart > 0:
            excedents.append((code, ecart, recu, cible))

    if not manquants and not excedents:
        return "Tous les formulaires mappés sont exactement à leur cible en effectifs."

    phrases = []
    if manquants:
        pire = max(manquants, key=lambda x: x[1])
        phrases.append(
            f"En volume, {pire[0]} est le plus éloigné de sa cible avec {pire[1]} "
            f"soumission(s) manquante(s) ({pire[2]}/{pire[3]})."
        )
    if excedents:
        pire = max(excedents, key=lambda x: x[1])
        phrases.append(
            f"À l'inverse, {pire[0]} dépasse sa cible de {pire[1]} soumission(s) "
            f"({pire[2]}/{pire[3]}), à recouper avec les anomalies en excès de la section 2."
        )
    return ' '.join(phrases)


def _superviseur_analysis_text(superviseur_table, form_codes):
    """Moyenne des taux sur le volet RDM (F01/F02/F07) PAR SUPERVISEUR —
    complète l'analyse par formulaire de la section 3 (qui regarde chaque
    formulaire indépendamment des personnes) avec une lecture par
    responsable de district, plus directement actionnable pour le suivi
    hiérarchique."""
    entries = []
    for s in superviseur_table.values():
        taux_vals = [s['forms'][c]['taux'] for c in form_codes
                     if s['forms'].get(c, {}).get('taux') is not None]
        if taux_vals:
            entries.append((s['nom'], round(sum(taux_vals) / len(taux_vals), 1)))
    if not entries:
        return "Aucune donnée exploitable pour comparer les superviseurs entre eux."

    entries.sort(key=lambda x: x[1])
    pire, meilleur = entries[0], entries[-1]
    n_faibles = sum(1 for _, t in entries if t < 60)
    phrase = (
        f"Sur le volet RDM ({'/'.join(form_codes)}), {meilleur[0]} affiche la moyenne la plus "
        f"élevée ({meilleur[1]} %), tandis que {pire[0]} est en retrait ({pire[1]} %)."
    )
    if n_faibles:
        phrase += f" {n_faibles} superviseur(s) sur {len(entries)} ont une moyenne sous les 60 %."
    return phrase


def _enqueteur_analysis_text(enqueteur_table, form_codes):
    """Analyse de la distribution des taux par enquêteur (complète le
    boxplot) — combien sont à 0 % (à relancer) et combien en net excédent
    (≥200 %, doublons potentiels) par formulaire, plutôt qu'une énumération
    nominative déjà couverte par le tableau qui suit."""
    phrases = []
    for code in form_codes:
        cells = [e['forms'].get(code, {}) for e in enqueteur_table.values()]
        n_total = sum(1 for c in cells if c.get('statut') not in (None, 'inconnu'))
        if not n_total:
            continue
        zero = sum(1 for c in cells if c.get('statut') == 'zero')
        excess = sum(1 for c in cells if c.get('statut') == 'verifier')
        detail = []
        if zero:
            detail.append(f"{zero} à 0 %")
        if excess:
            detail.append(f"{excess} en net excédent (≥200 %, doublons potentiels)")
        if detail:
            phrases.append(f"{code} : {', '.join(detail)} sur {n_total} enquêteur(s).")
    if not phrases:
        return ("Aucun enquêteur en situation extrême (0 % ou net excédent) détecté sur les "
                "formulaires observés.")
    return ' '.join(phrases)


def build_docx(cached, ref, computed_at=None):
    """Construit le rapport Word et renvoie les octets du fichier (.docx)."""
    national = cached.get('national', {})
    district_table = cached.get('district', {})
    anomalies_zero = cached.get('anomalies_zero', [])
    anomalies_excess = cached.get('anomalies_excess', [])
    form_codes = [c for c in rd.FORM_CODES if national.get(c, {}).get('statut') != 'inconnu']

    doc = Document()
    doc.styles['Normal'].font.size = Pt(10)

    annee = datetime.now().year
    title = doc.add_paragraph(f'SUIVI DE LA COMPLÉTUDE DE LA COLLECTE — SPAD {annee}')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.bold = True
    title.runs[0].font.size = Pt(18)
    title.runs[0].font.color.rgb = PRIMARY

    subtitle = doc.add_paragraph(
        f"Rapport généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
        + (f" — dernier calcul de complétude : {computed_at}" if computed_at else '')
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.italic = True
    subtitle.runs[0].font.size = Pt(9)
    subtitle.runs[0].font.color.rgb = RGBColor(0x6C, 0x6C, 0x6C)

    doc.add_paragraph()

    # ── 1. Indicateurs clés ──────────────────────────────────────────────
    h = doc.add_heading('1. Indicateurs clés', level=1)
    h.runs[0].font.color.rgb = PRIMARY
    recu_total, cible_total, taux_global, n_actifs = _national_totals(national, rd.FORM_CODES)
    n_etab = len(ref['etablissements'])
    n_district = len(ref['districts'])
    n_menages = 15 * n_etab
    _add_table(
        doc,
        ['Taux global', 'Soumissions Kobo', 'Formulaires actifs', 'Districts pilotes'],
        [[
            f"{taux_global} %" if taux_global is not None else '—',
            f"{recu_total:,}".replace(',', ' '),
            f"{n_actifs} / {len(rd.FORM_CODES)}",
            f"{n_district} ({n_etab} étab. · {n_menages} ménages)",
        ]],
    )
    doc.add_paragraph()

    # ── 2. Points d'alerte ───────────────────────────────────────────────
    h = doc.add_heading("2. Points d'alerte", level=1)
    h.runs[0].font.color.rgb = PRIMARY
    unites_zero = sorted(set(a['unite'] for a in anomalies_zero))
    if unites_zero:
        doc.add_paragraph(
            f"{len(unites_zero)} établissement(s)/district(s) affichent 0 % de soumission "
            f"sur au moins un formulaire déployé."
        )
    excess_by_form = {}
    for a in anomalies_excess:
        excess_by_form.setdefault(a['formulaire_label'], []).append(a)
    if excess_by_form:
        detail = ', '.join(f"{label} ({len(items)})" for label, items in excess_by_form.items())
        doc.add_paragraph(
            f"{len(anomalies_excess)} soumission(s) dépassent la cible attendue, "
            f"doublons potentiels à investiguer : {detail}."
        )
    if not unites_zero and not excess_by_form:
        doc.add_paragraph("Aucune anomalie détectée sur les formulaires mappés à ce jour.")
    doc.add_paragraph()

    # ── 3. Complétude par formulaire ─────────────────────────────────────
    h = doc.add_heading('3. Complétude par formulaire', level=1)
    h.runs[0].font.color.rgb = PRIMARY

    anomalies_par_formulaire = {}
    for a in anomalies_zero + anomalies_excess:
        anomalies_par_formulaire[a['formulaire']] = anomalies_par_formulaire.get(a['formulaire'], 0) + 1

    history = tendance.load_history(days=30)
    velocite = _velocity_and_projection(national, rd.FORM_CODES, history)

    rows = []
    for code in rd.FORM_CODES:
        r = national.get(code, {})
        statut = r.get('statut', 'inconnu')
        cible, recu = r.get('cible'), r.get('recu', 0)
        taux = f"{r['taux']} %" if r.get('taux') is not None else '—'
        if statut == 'inconnu':
            ecart = '—'
        elif cible and recu >= cible:
            ecart = f"+{recu - cible}" if recu > cible else '0'
        elif cible:
            ecart = f"−{cible - recu}"
        else:
            ecart = '—'
        n_anom = anomalies_par_formulaire.get(code, 0)
        dr = cp.district_reel(district_table, code) if district_table else None
        cible_reelle = f"{dr['pct']} % ({dr['atteints']}/{dr['total']})" if dr else '—'
        rows.append([
            f"{code} — {rd.FORM_LABELS[code]}",
            cible if cible is not None else '—',
            recu if statut != 'inconnu' else '—',
            ecart, taux, STATUT_LABEL.get(statut, statut),
            str(n_anom) if n_anom else '—',
            cible_reelle,
        ])
    _add_table(
        doc,
        ['Formulaire', 'Cible', 'Reçu', 'Écart', 'Taux', 'Statut', 'Anomalies', 'Cible réelle atteinte'],
        rows,
    )
    doc.add_paragraph()

    def _add_chart_with_analysis(chart_path, caption, analysis_text, width=6.2):
        """Insère un graphique suivi immédiatement d'un paragraphe d'analyse
        — pour que chaque visuel de ce rapport soit systématiquement
        accompagné d'une lecture en prose, pas seulement d'une légende.
        L'analyse est un texte dérivé des données (pas du rendu de l'image) :
        elle est ajoutée même si le graphique échoue à se générer (ex.
        Kaleido indisponible), pour ne jamais perdre l'information faute
        d'un visuel manquant."""
        if chart_path:
            try:
                doc.add_picture(chart_path, width=Inches(width))
                cap = doc.add_paragraph(caption)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].font.size = Pt(8)
                cap.runs[0].font.italic = True
                cap.runs[0].font.color.rgb = RGBColor(0x6C, 0x6C, 0x6C)
            finally:
                try:
                    os.remove(chart_path)
                except OSError:
                    pass
        if analysis_text:
            p = doc.add_paragraph()
            run = p.add_run('Analyse : ')
            run.font.bold = True
            p.add_run(analysis_text)
        doc.add_paragraph()

    _add_chart_with_analysis(
        _summary_chart_image(national, rd.FORM_CODES),
        'Taux de complétude par formulaire',
        _form_rate_analysis_text(national, rd.FORM_CODES, anomalies_par_formulaire, district_table),
    )
    _add_chart_with_analysis(
        _form_counts_chart_image(national, rd.FORM_CODES),
        'Soumissions reçues vs cible, en effectifs bruts',
        _form_counts_analysis_text(national, rd.FORM_CODES),
    )

    # ── 3bis. Dynamique de collecte (lecture suivi-évaluation) ───────────
    doc.add_heading('Dynamique de collecte et projection', level=2)
    en_cours = [c for c in rd.FORM_CODES if national.get(c, {}).get('statut') == 'en_cours']
    if not en_cours:
        doc.add_paragraph(
            "Aucun formulaire en cours de collecte à analyser (soit non mappé, soit déjà à la "
            "cible ou au-delà)."
        )
    else:
        any_disponible = False
        for code in en_cours:
            v = velocite.get(code, {'disponible': False})
            r = national.get(code, {})
            n_anom = anomalies_par_formulaire.get(code, 0)
            p = doc.add_paragraph()
            run = p.add_run(f"{code} — {rd.FORM_LABELS[code]} : ")
            run.font.bold = True
            if not v['disponible']:
                p.add_run(
                    "historique insuffisant (moins de 2 jours de calcul distincts) pour estimer "
                    "une vélocité de collecte — se reconstituera au fil des prochains calculs."
                )
                continue
            any_disponible = True
            vitesse = v['vitesse_jour']
            reste = max(0, (r.get('cible') or 0) - (r.get('recu') or 0))
            if vitesse > 0:
                p.add_run(
                    f"progresse à ~{vitesse} soumission(s)/jour en moyenne sur les "
                    f"{v['n_jours_observes']} dernier(s) jour(s) observé(s). "
                )
                if v['projection'] == 'estimee':
                    p.add_run(
                        f"Au rythme actuel, la cible ({reste} soumission(s) restante(s)) serait "
                        f"atteinte dans environ {v['jours_restants']} jour(s)."
                    )
            elif vitesse == 0:
                p.add_run(
                    f"AUCUNE progression détectée sur les {v['n_jours_observes']} dernier(s) "
                    f"jour(s) observé(s) — collecte possiblement à l'arrêt, "
                    f"{reste} soumission(s) restant(s) à obtenir. À investiguer auprès des "
                    f"superviseurs concernés."
                )
            else:
                p.add_run(
                    f"taux en RECUL sur les {v['n_jours_observes']} dernier(s) jour(s) "
                    f"({vitesse}/jour) — signe possible d'une correction de doublons ou d'un "
                    f"problème de synchronisation Kobo à vérifier."
                )
            if n_anom:
                p.add_run(f" {n_anom} anomalie(s) associée(s) (voir section 2).")
        if not any_disponible:
            doc.add_paragraph(
                "Aucune estimation de vélocité disponible pour l'instant — au moins 2 calculs de "
                "complétude sur des jours calendaires distincts sont nécessaires (un point "
                "d'historique par jour, voir modules/tendance.py).",
            ).runs[0].font.italic = True
    doc.add_paragraph()

    # ── 4. Complétude par district ───────────────────────────────────────
    n_district_actifs = sum(
        1 for d in district_table.values()
        if any(d['forms'].get(c, {}).get('statut') not in (None, 'inconnu') for c in form_codes)
    )
    h = doc.add_heading(f'4. Complétude par district ({n_district_actifs}/{len(district_table)})', level=1)
    h.runs[0].font.color.rgb = PRIMARY
    headers = ['District'] + [f"{c}" for c in form_codes]
    rows = []
    for dcode in sorted(district_table, key=lambda c: district_table[c]['nom']):
        d = district_table[dcode]
        row = [d['nom']]
        for c in form_codes:
            cell = d['forms'].get(c, {})
            row.append(f"{cell['taux']} %" if cell.get('taux') is not None else '—')
        rows.append(row)
    _add_table(doc, headers, rows)
    doc.add_paragraph()
    district_chart = _grouped_bar_chart(district_table, form_codes, 'Taux de complétude par district et par formulaire')
    if district_chart:
        try:
            doc.add_picture(district_chart, width=Inches(6.2))
        finally:
            try:
                os.remove(district_chart)
            except OSError:
                pass
    doc.add_paragraph()
    for bullet in _district_form_bullets(district_table, form_codes):
        doc.add_paragraph(bullet, style='List Bullet')
    doc.add_paragraph()

    # ── 5. Complétude par superviseur ────────────────────────────────────
    superviseur_table = cached.get('superviseur', {})
    sup_form_codes = [c for c in cp.SUPERVISEUR_FORMS if c in form_codes]
    if superviseur_table and sup_form_codes:
        h = doc.add_heading(f'5. Complétude par superviseur ({len(superviseur_table)})', level=1)
        h.runs[0].font.color.rgb = PRIMARY
        headers = ['Superviseur', 'District'] + sup_form_codes
        rows = []
        for code, s in sorted(superviseur_table.items(), key=lambda kv: kv[1]['nom']):
            row = [s['nom'], s.get('sous_titre', '—')]
            for c in sup_form_codes:
                cell = s['forms'].get(c, {})
                row.append(f"{cell['taux']} %" if cell.get('taux') is not None else '—')
            rows.append(row)
        _add_table(doc, headers, rows)
        doc.add_paragraph()
        _add_chart_with_analysis(
            _grouped_bar_chart(superviseur_table, sup_form_codes, 'Taux de complétude par superviseur'),
            'Taux de complétude par superviseur',
            _superviseur_analysis_text(superviseur_table, sup_form_codes),
        )

    # ── 6. Complétude par enquêteur ──────────────────────────────────────
    enqueteur_table = cached.get('enqueteur', {})
    enq_form_codes = [c for c in cp.ENQUETEUR_FORMS if c in form_codes]
    if enqueteur_table and enq_form_codes:
        h = doc.add_heading(f'6. Complétude par enquêteur ({len(enqueteur_table)})', level=1)
        h.runs[0].font.color.rgb = PRIMARY
        doc.add_paragraph(
            "Distribution des taux (tous les enquêteurs) puis détail nominatif — un graphique en "
            "barres par enquêteur ne serait pas lisible au-delà de quelques dizaines d'unités."
        )
        _add_chart_with_analysis(
            _distribution_chart(enqueteur_table, enq_form_codes, 'Distribution des taux de complétude par enquêteur'),
            'Distribution des taux de complétude par enquêteur',
            _enqueteur_analysis_text(enqueteur_table, enq_form_codes),
            width=5.5,
        )
        headers = ['Enquêteur', 'District'] + enq_form_codes
        rows = []
        for code, e in sorted(enqueteur_table.items(), key=lambda kv: kv[1]['nom']):
            row = [e['nom'], e.get('sous_titre', '—')]
            for c in enq_form_codes:
                cell = e['forms'].get(c, {})
                row.append(f"{cell['taux']} %" if cell.get('taux') is not None else '—')
            rows.append(row)
        _add_table(doc, headers, rows)
        doc.add_paragraph()

    # ── 7. Actions suggérées ─────────────────────────────────────────────
    h = doc.add_heading('7. Actions suggérées', level=1)
    h.runs[0].font.color.rgb = PRIMARY
    actions = []
    if anomalies_excess:
        actions.append(
            f"Faire vérifier par les superviseurs les {len(anomalies_excess)} soumission(s) "
            f"excédentaire(s) (doublons potentiels) avant tout nettoyage de la base."
        )
    if unites_zero:
        actions.append(
            f"Relancer les {len(unites_zero)} établissement(s)/district(s) n'ayant soumis "
            f"aucune donnée sur au moins un formulaire déployé."
        )
    retard = [f"{c} — {rd.FORM_LABELS[c]}" for c in form_codes
              if (national.get(c, {}).get('taux') or 0) < 60]
    if retard:
        actions.append("Accélérer la collecte sur les formulaires les plus en retard (< 60 %) : "
                        + ', '.join(retard) + '.')
    if not actions:
        actions.append("Aucune action corrective urgente identifiée sur les données actuelles.")
    for a in actions:
        doc.add_paragraph(a, style='List Bullet')

    doc.add_paragraph()
    note = doc.add_paragraph(
        "Rapport généré automatiquement par SPAD PHAKTS Analyzer à partir du dernier calcul "
        "de complétude. Les actions ci-dessus sont dérivées des anomalies détectées ; les "
        "décisions de réunion (responsables, budget, calendrier) restent à compléter manuellement."
    )
    note.runs[0].font.size = Pt(8)
    note.runs[0].font.italic = True
    note.runs[0].font.color.rgb = RGBColor(0x6C, 0x6C, 0x6C)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
