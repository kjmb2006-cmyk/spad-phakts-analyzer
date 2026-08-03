#!/usr/bin/env python3
"""
Test script — rapport de complétude Word (modules/completude_report.py).

Vérifie la structure du document généré (titres, tableaux) à partir de
données de complétude réalistes, et le câblage de la route Flask
/completude/export.docx (avec et sans cache disponible).
"""
import os
import json
import datetime
import pandas as pd
import docx

from modules import reference_data as rd, completeness as cp, completude_report as rep, tendance

print("=" * 70)
print("TEST — Rapport de complétude Word (modules/completude_report.py)")
print("=" * 70)

ref = rd.load()
etabs = list(ref['etablissements'].values())


def _make_df(ratio_fn, code):
    field = cp.FORM_FIELDS[code]['etablissement']  # varie selon le formulaire (ex. F02 = 'F02_01__E')
    rows = []
    for i, e in enumerate(etabs):
        cible = rd.target_for(ref, code, etablissement_code=e['code'])
        n = max(0, int(round((cible or 1) * ratio_fn(i))))
        rows.extend([{field: e['code']} for _ in range(n)])
    return pd.DataFrame(rows)


form_dataframes = {
    'F5': _make_df(lambda i: 0.5 + (i % 5) * 0.05, 'F5'),
    'F7': _make_df(lambda i: 0.0 if i % 20 == 0 else 0.95, 'F7'),  # quelques 0 % -> alertes
    'F8': _make_df(lambda i: 1.1, 'F8'),                            # dépassement -> anomalie
    'F02': _make_df(lambda i: 1.0, 'F02'),                          # formulaire superviseur mappé
}

national = cp.national_summary(ref, form_dataframes)
cached = {
    'national':         national,
    'district':         cp.district_table(ref, form_dataframes),
    'region':           cp.region_table(ref, form_dataframes),
    'enqueteur':        cp.enqueteur_table(ref, form_dataframes),
    'superviseur':      cp.superviseur_table(ref, form_dataframes),
    'anomalies_zero':   cp.all_anomalies_zero(ref, form_dataframes),
    'anomalies_excess': cp.all_anomalies_excess(ref, form_dataframes),
    'export':           cp.export_rows(ref, form_dataframes),
}
assert len(cached['anomalies_zero']) > 0, "le jeu de test doit produire au moins une alerte à 0 %"
print(f"OK — jeu de données simulé : {len(cached['anomalies_zero'])} alerte(s) 0 %,"
      f" {len(cached['anomalies_excess'])} excédent(s)")
# national['F5'] = en_cours, taux=60.0 ; national['F7'] = en_cours, taux=88.7 (valeurs stables,
# vérifiées au préalable) — utilisées ci-dessous pour simuler un historique cohérent.

# ── Historique simulé : F5 progresse (50 % -> 60 %), F7 stagne (88.7 % fixe) ──
# Isole le fichier d'historique de test de celui de l'utilisateur.
tendance.HISTORY_PATH = '/tmp/spad_test_historique_rapport.jsonl'
if os.path.exists(tendance.HISTORY_PATH):
    os.remove(tendance.HISTORY_PATH)
today = datetime.date.today()
history_points = [
    {'date': (today - datetime.timedelta(days=2)).isoformat(), 'taux': {'F5': 50.0, 'F7': 88.7}},
    {'date': (today - datetime.timedelta(days=1)).isoformat(), 'taux': {'F5': 55.0, 'F7': 88.7}},
]
with open(tendance.HISTORY_PATH, 'w', encoding='utf-8') as f:
    for p in history_points:
        f.write(json.dumps(p, ensure_ascii=False) + '\n')

docx_bytes = rep.build_docx(cached, ref, computed_at="31/07/2026 22:00:00")
assert len(docx_bytes) > 5000
print(f"OK — document généré ({len(docx_bytes)} octets)")

tmp_path = '/tmp/spad_test_rapport.docx'
with open(tmp_path, 'wb') as f:
    f.write(docx_bytes)
d = docx.Document(tmp_path)
os.remove(tmp_path)

assert 'SUIVI DE LA COMPLÉTUDE' in d.paragraphs[0].text
headings = [p.text for p in d.paragraphs if p.style.name.startswith('Heading')]
assert any('Indicateurs clés' in h for h in headings)
assert any('alerte' in h for h in headings)
assert any('formulaire' in h.lower() for h in headings)
assert any('district' in h.lower() for h in headings)
assert any('superviseur' in h.lower() for h in headings)
assert any('enquêteur' in h.lower() for h in headings)
assert any('Actions' in h for h in headings)
assert len(d.tables) == 5, (
    f"attendu 5 tableaux (indicateurs / formulaires / districts / superviseurs / enquêteurs), "
    f"trouvé {len(d.tables)}")
print("OK — structure conforme : titre, 7 sections, 5 tableaux (dont superviseur/enquêteur)")

# Graphique de synthèse + par district + par superviseur (grouped bar) +
# distribution par enquêteur (boxplot) = 4 images attendues.
assert len(d.inline_shapes) >= 4, (
    f"attendu au moins 4 images (synthèse, district, superviseur, distribution enquêteur), "
    f"trouvé {len(d.inline_shapes)}")
print(f"OK — graphiques intégrés ({len(d.inline_shapes)} image(s) dans le document)")

# Le tableau des formulaires doit lister les 7 formulaires officiels, même
# ceux non mappés (statut 'Non calculé'), avec les colonnes Écart, Anomalies
# et Cible réelle atteinte (lecture suivi-évaluation).
form_table = d.tables[1]
header = [c.text for c in form_table.rows[0].cells]
assert header == ['Formulaire', 'Cible', 'Reçu', 'Écart', 'Taux', 'Statut', 'Anomalies', 'Cible réelle atteinte'], header
form_col0 = [row.cells[0].text for row in form_table.rows[1:]]
assert len(form_col0) == 7
assert any('F6' in c for c in form_col0) and any('Non calculé' in row.cells[5].text
                                                    for row in form_table.rows[1:] if 'F6' in row.cells[0].text)
print("OK — les 7 formulaires officiels apparaissent, avec les colonnes Écart/Anomalies/Cible réelle atteinte")

f5_row = next(row for row in form_table.rows[1:] if row.cells[0].text.startswith('F5'))
assert f5_row.cells[3].text == '−720', f5_row.cells[3].text  # écart = 1800 - 1080
f7_row = next(row for row in form_table.rows[1:] if row.cells[0].text.startswith('F7'))
assert f7_row.cells[6].text != '—', "F7 doit avoir des anomalies associées (0 % simulés)"
print("OK — écart à la cible et anomalies par formulaire correctement calculés")

# ── Dynamique de collecte : F5 progresse (50%->55% en 1 jour = +90 soum./j,
# reste 720 -> ~8 jours), F7 stagne (88.7% fixe) ──
full_text = '\n'.join(p.text for p in d.paragraphs)
assert 'Dynamique de collecte' in full_text
assert 'F5 — Tabac' in full_text and 'progresse à ~90.0 soumission' in full_text
assert 'atteinte dans environ 8 jour' in full_text
assert 'AUCUNE progression détectée' in full_text  # F7, taux stable à 88.7 % sur l'historique simulé
print("OK — vélocité de collecte estimée depuis l'historique (F5 en progression, F7 stagnant)")

os.remove(tendance.HISTORY_PATH)

# ── Commentaires « par district » : un commentaire PAR FORMULAIRE (pas par
# volet agrégé), avec le district le plus/moins avancé et l'écart restant.
assert 'district le plus avancé' in full_text and 'le moins avancé' in full_text
assert 'F5 — Tabac' in full_text.split('Dynamique de collecte')[1]  # présent aussi côté section district
print("OK — commentaires « par district » faits par formulaire (min/max/écart), pas par volet agrégé")

# ── Tableaux superviseur / enquêteur (seuls les formulaires effectivement
# mappés apparaissent en colonne : F02 côté superviseur, F5/F7/F8 côté
# enquêteur — F01/F07/F6 non mappés dans ce jeu de test) ──
sup_table = next(t for t in d.tables if t.rows[0].cells[0].text == 'Superviseur')
assert [c.text for c in sup_table.rows[0].cells] == ['Superviseur', 'District', 'F02']
assert len(sup_table.rows) - 1 == 12, f"attendu 12 superviseurs, trouvé {len(sup_table.rows) - 1}"
print("OK — tableau « par superviseur » présent (12 lignes, F02 mappé)")

enq_table = next(t for t in d.tables if t.rows[0].cells[0].text == 'Enquêteur')
assert [c.text for c in enq_table.rows[0].cells] == ['Enquêteur', 'District', 'F5', 'F7', 'F8']
assert len(enq_table.rows) - 1 == 60, f"attendu 60 enquêteurs, trouvé {len(enq_table.rows) - 1}"
print("OK — tableau « par enquêteur » présent (60 lignes, F5/F7/F8 mappés)")

print()
print("=" * 70)
print("TESTS STRUCTURELS PASSÉS — vérification de la route Flask")
print("=" * 70)

from app import app  # noqa: E402 (après configuration du chemin par les tests précédents)

client = app.test_client()
fname = os.path.join(app.config['UPLOAD_FOLDER'], 'completude_test_report.json')
with open(fname, 'w', encoding='utf-8') as f:
    json.dump(cached, f)

with client.session_transaction() as sess:
    sess['completude_path'] = fname
    sess['completude_computed_at'] = '31/07/2026 22:00:00'

r = client.get('/completude/export.docx')
assert r.status_code == 200
assert 'wordprocessingml' in r.mimetype
assert len(r.data) > 5000
print("OK — route /completude/export.docx : 200, document Word valide")

os.remove(fname)
with client.session_transaction() as sess:
    sess.pop('completude_path', None)
r2 = client.get('/completude/export.docx', follow_redirects=True)
assert r2.status_code == 200
assert "Calculez d'abord" in r2.get_data(as_text=True)
print("OK — sans cache disponible : redirection propre (pas d'erreur 500)")

print()
print("=" * 70)
print("TOUS LES TESTS RAPPORT WORD SONT PASSÉS")
print("=" * 70)
